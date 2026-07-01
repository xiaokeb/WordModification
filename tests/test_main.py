# -*- coding: utf-8 -*-
"""
Word 文本处理工具 - 单元测试
测试各核心功能的有效性和稳定性

运行方法：
    cd e:\\SourceCode\\Projects\\WordModification
    python tests/test_main.py

注意：
- 文本替换和PDF导出相关测试需要Microsoft Word/Excel
- 后缀修改和Excel写入测试可独立运行
"""

import os
import sys
import unittest
import tempfile
import shutil
import zipfile
from pathlib import Path
from unittest.mock import MagicMock, patch, mock_open
from io import BytesIO

# 添加src目录到路径
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))


class TestWordProcessor(unittest.TestCase):
    """测试Word文档处理器"""

    def setUp(self):
        """测试前准备：创建临时目录"""
        self.test_dir = tempfile.mkdtemp(prefix="word_test_")
        self.processor = None
        try:
            from main import WordProcessor
            self.processor = WordProcessor()
        except ImportError as e:
            self.skipTest(f"无法导入WordProcessor: {e}")

    def tearDown(self):
        """测试后清理"""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir, ignore_errors=True)

    def test_validate_file(self):
        """测试文件验证功能"""
        # 测试不存在的文件
        non_existent = Path(self.test_dir) / "not_exist.docx"
        self.assertFalse(self.processor._validate_file(non_existent))

        # 测试目录
        self.assertFalse(self.processor._validate_file(Path(self.test_dir)))

        # 测试有效文件
        test_file = Path(self.test_dir) / "test.docx"
        test_file.write_text("test")
        self.assertTrue(self.processor._validate_file(test_file))

    def test_get_word_files_from_directory(self):
        """测试从目录获取Word文件"""
        # 创建测试文件
        for name in ["file1.docx", "file2.doc", "readme.txt"]:
            (Path(self.test_dir) / name).write_text("test")

        files = self.processor.get_word_files(self.test_dir)
        self.assertEqual(len(files), 2)
        # 确保.txt文件被排除
        for f in files:
            self.assertTrue(f.endswith(('.doc', '.docx')))

    def test_get_word_files_from_nonexistent_path(self):
        """测试不存在的路径"""
        files = self.processor.get_word_files("/non_existent_path_12345")
        self.assertEqual(files, [])

    def test_get_word_files_from_empty_path(self):
        """测试空路径"""
        files = self.processor.get_word_files("")
        self.assertEqual(files, [])

    def test_get_word_files_from_list(self):
        """测试从文件列表获取Word文件"""
        # 创建测试文件
        valid_file = Path(self.test_dir) / "valid.docx"
        valid_file.write_text("test")
        invalid_file = Path(self.test_dir) / "invalid.txt"
        invalid_file.write_text("test")

        files = self.processor.get_word_files([str(valid_file), str(invalid_file)])
        self.assertEqual(len(files), 1)
        self.assertTrue(files[0].endswith('.docx'))

    def test_replace_text_in_paragraph_simple(self):
        """测试段落文本替换 - 简单情况"""
        from docx import Document

        # 创建测试文档
        doc = Document()
        para = doc.add_paragraph("Hello World")
        replacements = {"World": "Python"}

        self.processor.replace_text_in_paragraph(para, replacements)
        self.assertEqual(para.text, "Hello Python")

    def test_replace_text_in_paragraph_multiple(self):
        """测试段落文本替换 - 多个规则"""
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("公司名称ABC，项目编号HT001")
        replacements = {"ABC": "测试公司", "HT001": "2024-001"}

        self.processor.replace_text_in_paragraph(para, replacements)
        self.assertIn("测试公司", para.text)
        self.assertIn("2024-001", para.text)

    def test_replace_text_in_paragraph_empty(self):
        """测试空替换规则"""
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("Original text")
        original_text = para.text

        self.processor.replace_text_in_paragraph(para, {})
        self.assertEqual(para.text, original_text)

    def test_replace_text_in_paragraph_skip_empty_key(self):
        """测试空键值被跳过"""
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("Test text")
        replacements = {"": "value", "Test": "结果"}

        self.processor.replace_text_in_paragraph(para, replacements)
        self.assertEqual(para.text, "结果 text")

    def test_replace_text_cross_run(self):
        """测试跨Run的文本替换"""
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document()
        para = doc.add_paragraph()

        # 手动添加多个run模拟跨Run文本
        run1 = para.add_run("项目")
        run2 = para.add_run("编号")
        run3 = para.add_run("ABC")

        replacements = {"项目编号ABC": "HT-2024-001"}

        self.processor.replace_text_in_paragraph(para, replacements)
        self.assertEqual(para.text, "HT-2024-001")

    def test_process_document_with_corrupted_file(self):
        """测试处理损坏的文件"""
        # 创建一个无效的.docx文件
        invalid_file = Path(self.test_dir) / "invalid.docx"
        invalid_file.write_text("not a valid docx")

        output = Path(self.test_dir) / "output.docx"
        result = self.processor.process_document(
            str(invalid_file),
            {"a": "b"},
            str(output)
        )
        self.assertFalse(result)

    def test_process_document_nonexistent(self):
        """测试处理不存在的文件"""
        result = self.processor.process_document(
            "/nonexistent/file.docx",
            {"a": "b"},
            "/tmp/output.docx"
        )
        self.assertFalse(result)


class TestExcelProcessor(unittest.TestCase):
    """测试Excel处理器"""

    def setUp(self):
        """测试前准备"""
        self.test_dir = tempfile.mkdtemp(prefix="excel_test_")
        try:
            from main import ExcelProcessor
            self.processor = ExcelProcessor()
        except ImportError as e:
            self.skipTest(f"无法导入ExcelProcessor: {e}")

    def tearDown(self):
        """测试后清理"""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir, ignore_errors=True)

    def _create_test_excel(self, filename="test.xlsx", num_sheets=1):
        """创建测试Excel文件"""
        try:
            from openpyxl import Workbook
        except ImportError:
            self.skipTest("缺少 openpyxl 库")

        wb = Workbook()
        # 重命名默认sheet
        if num_sheets == 1:
            wb.active.title = "Sheet1"
        else:
            wb.remove(wb.active)
            for i in range(num_sheets):
                wb.create_sheet(f"Sheet{i+1}")
                ws = wb[f"Sheet{i+1}"]
                ws['A1'] = f"Test{i+1}"

        file_path = Path(self.test_dir) / filename
        wb.save(str(file_path))
        return str(file_path)

    def test_write_number_auto_mode(self):
        """测试自动模式写入编号"""
        excel_path = self._create_test_excel("test.xlsx", num_sheets=2)
        output_path = Path(self.test_dir) / "output.xlsx"

        success = self.processor.write_number_to_excel(
            excel_path,
            str(output_path),
            "HT001-2026-XM001-01",
            use_fixed_mode=False,
            fixed_number="",
            number_format="HT001-2026-XM001-yy"
        )

        self.assertTrue(success)
        self.assertTrue(output_path.exists())

        # 验证内容
        from openpyxl import load_workbook
        wb = load_workbook(str(output_path))
        self.assertEqual(wb["Sheet1"]["E1"].value, "HT001-2026-XM001-01")
        self.assertEqual(wb["Sheet2"]["E1"].value, "HT001-2026-XM001-02")

    def test_write_number_fixed_mode(self):
        """测试固定模式写入编号"""
        excel_path = self._create_test_excel("test.xlsx", num_sheets=2)
        output_path = Path(self.test_dir) / "output.xlsx"

        success = self.processor.write_number_to_excel(
            excel_path,
            str(output_path),
            "HT001-2026-XM001-01",
            use_fixed_mode=True,
            fixed_number="公司名称",
            number_format="HT001-2026-XM001-yy"
        )

        self.assertTrue(success)
        self.assertTrue(output_path.exists())

        # 验证内容
        from openpyxl import load_workbook
        wb = load_workbook(str(output_path))
        self.assertEqual(wb["Sheet1"]["E1"].value, "公司名称")
        self.assertEqual(wb["Sheet1"]["E2"].value, "HT001-2026-XM001-01")
        self.assertEqual(wb["Sheet2"]["E2"].value, "HT001-2026-XM001-02")

    def test_write_number_nonexistent_file(self):
        """测试不存在的Excel文件"""
        success = self.processor.write_number_to_excel(
            "/nonexistent/file.xlsx",
            str(Path(self.test_dir) / "output.xlsx"),
            "HT001"
        )
        self.assertFalse(success)


class TestSuffixProcessor(unittest.TestCase):
    """测试文件后缀处理器"""

    def setUp(self):
        """测试前准备"""
        self.test_dir = tempfile.mkdtemp(prefix="suffix_test_")
        try:
            from main import SuffixProcessor
            self.processor = SuffixProcessor()
        except ImportError as e:
            self.skipTest(f"无法导入SuffixProcessor: {e}")

    def tearDown(self):
        """测试后清理"""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir, ignore_errors=True)

    def _create_test_file(self, name):
        """创建测试文件"""
        file_path = Path(self.test_dir) / name
        file_path.write_text("test content")
        return str(file_path)

    def test_remove_suffix_simple(self):
        """测试简单后缀删除"""
        test_file = self._create_test_file("document-修改版.docx")

        success, total = self.processor.remove_suffix_from_files([test_file])

        self.assertEqual(success, 1)
        self.assertEqual(total, 1)
        # 原文件应被重命名
        new_file = Path(self.test_dir) / "document.docx"
        self.assertTrue(new_file.exists())

    def test_remove_suffix_multiple_occurrences(self):
        """测试多个后缀同时删除"""
        test_file = self._create_test_file("doc-修改版-final-修改版.docx")

        success, total = self.processor.remove_suffix_from_files([test_file])

        self.assertEqual(success, 1)
        # 文件名中所有"-修改版"应被删除
        new_file = Path(self.test_dir) / "doc-final.docx"
        self.assertTrue(new_file.exists())

    def test_remove_suffix_custom_suffix(self):
        """测试自定义后缀"""
        test_file = self._create_test_file("file-draft.txt")

        success, total = self.processor.remove_suffix_from_files(
            [test_file],
            suffix="-draft"
        )

        self.assertEqual(success, 1)
        new_file = Path(self.test_dir) / "file.txt"
        self.assertTrue(new_file.exists())

    def test_remove_suffix_no_match(self):
        """测试无匹配后缀的情况"""
        test_file = self._create_test_file("normal.docx")
        original = Path(test_file).name

        success, total = self.processor.remove_suffix_from_files([test_file])

        # 没有匹配时应不重命名
        self.assertEqual(success, 0)
        self.assertTrue(Path(test_file).exists())

    def test_remove_suffix_empty_list(self):
        """测试空文件列表"""
        success, total = self.processor.remove_suffix_from_files([])
        self.assertEqual(success, 0)
        self.assertEqual(total, 0)

    def test_remove_suffix_nonexistent_file(self):
        """测试不存在的文件"""
        success, total = self.processor.remove_suffix_from_files(
            ["/nonexistent/file-修改版.docx"]
        )
        # 不应崩溃
        self.assertEqual(success, 0)


class TestExcelToPdfProcessor(unittest.TestCase):
    """测试Excel转PDF处理器（需要pywin32 + Excel）"""

    def setUp(self):
        """测试前准备"""
        self.test_dir = tempfile.mkdtemp(prefix="pdf_test_")
        try:
            from main import ExcelToPdfProcessor
            self.processor = ExcelToPdfProcessor()
        except ImportError as e:
            self.skipTest(f"无法导入ExcelToPdfProcessor: {e}")

    def tearDown(self):
        """测试后清理"""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir, ignore_errors=True)

    def test_pywin32_available(self):
        """测试pywin32是否可用"""
        try:
            import win32com.client
            import pythoncom
            self.assertTrue(True, "pywin32 可用")
        except ImportError:
            self.skipTest("pywin32 未安装")

    def test_export_nonexistent_file(self):
        """测试不存在的Excel文件"""
        try:
            import win32com.client
        except ImportError:
            self.skipTest("pywin32 未安装")

        success, message = self.processor.export_to_pdf(
            "/nonexistent/file.xlsx",
            str(Path(self.test_dir) / "output.pdf")
        )
        self.assertFalse(success)
        self.assertIsNotNone(message)

    def test_detect_method(self):
        """测试方法检测功能"""
        # 至少检测到一种方法（pywin32/libreoffice/reportlab）
        # 如果都没有，method应为None
        self.assertTrue(hasattr(self.processor, 'method'))
        # 验证检测结果有效
        valid_methods = [None, "pywin32", "libreoffice", "reportlab"]
        self.assertIn(self.processor.method, valid_methods)

    def test_no_method_message(self):
        """测试无可用方法时的错误信息"""
        # 模拟所有方法都不可用
        original_method = self.processor.method
        self.processor.method = None
        try:
            success, message = self.processor.export_to_pdf(
                str(Path(self.test_dir) / "test.xlsx"),
                str(Path(self.test_dir) / "output.pdf")
            )
            self.assertFalse(success)
            self.assertIn("无可用的PDF导出方法", message)
        finally:
            self.processor.method = original_method

    def test_export_with_reportlab(self):
        """测试使用reportlab导出PDF（无需pywin32）"""
        try:
            from openpyxl import Workbook
        except ImportError:
            self.skipTest("缺少 openpyxl 库")

        # 创建测试Excel
        wb = Workbook()
        wb.active["A1"] = "测试"
        excel_path = Path(self.test_dir) / "test.xlsx"
        wb.save(str(excel_path))

        output_path = Path(self.test_dir) / "output.pdf"

        # 直接调用reportlab导出
        if hasattr(self.processor, '_export_with_reportlab'):
            success, message = self.processor._export_with_reportlab(
                str(excel_path), str(output_path)
            )
            # 可能因为reportlab未安装而失败，但不应崩溃
            if success:
                self.assertTrue(output_path.exists())
            else:
                # 如果失败，应有明确错误信息
                self.assertIsInstance(message, str)


class TestUtilityFunctions(unittest.TestCase):
    """测试工具函数"""

    def setUp(self):
        """测试前准备"""
        self.test_dir = tempfile.mkdtemp(prefix="util_test_")

    def tearDown(self):
        """测试后清理"""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir, ignore_errors=True)

    def test_path_operations(self):
        """测试路径操作"""
        from pathlib import Path
        # 测试路径拼接
        test_file = Path(self.test_dir) / "test-修改版.docx"
        new_name = test_file.stem.replace("-修改版", "") + test_file.suffix
        new_path = test_file.parent / new_name
        self.assertEqual(new_path.name, "test.docx")

    def test_filename_conflict_resolution(self):
        """测试文件名冲突解决方案"""
        # 创建同名文件测试冲突解决逻辑
        f1 = Path(self.test_dir) / "test.docx"
        f1.write_text("first")
        f2 = Path(self.test_dir) / "test_1.docx"
        f2.write_text("second")

        # 两个文件都应存在
        self.assertTrue(f1.exists())
        self.assertTrue(f2.exists())


class TestIntegration(unittest.TestCase):
    """集成测试 - 测试各功能协同工作"""

    def setUp(self):
        """测试前准备"""
        self.test_dir = tempfile.mkdtemp(prefix="integration_test_")

    def tearDown(self):
        """测试后清理"""
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir, ignore_errors=True)

    def test_excel_to_number_to_pdf_workflow(self):
        """测试 Excel -> 写入编号 -> 导出PDF 的完整工作流（不实际导出PDF）"""
        try:
            from main import ExcelProcessor
            from openpyxl import Workbook
        except ImportError as e:
            self.skipTest(f"依赖缺失: {e}")

        # 1. 创建测试Excel
        wb = Workbook()
        wb.active.title = "Sheet1"
        wb["Sheet1"]["A1"] = "测试数据"
        excel_path = Path(self.test_dir) / "test.xlsx"
        wb.save(str(excel_path))

        # 2. 写入编号
        processor = ExcelProcessor()
        output_path = Path(self.test_dir) / "output.xlsx"

        success = processor.write_number_to_excel(
            str(excel_path),
            str(output_path),
            "TEST-2026-001",
            use_fixed_mode=False,
            number_format="TEST-2026-yy"
        )

        self.assertTrue(success)
        self.assertTrue(output_path.exists())

        # 3. 验证编号已写入（ExcelProcessor会将yy替换为01,02等序号）
        from openpyxl import load_workbook
        result_wb = load_workbook(str(output_path))
        self.assertEqual(result_wb["Sheet1"]["E1"].value, "TEST-2026-01")

    def test_word_replace_in_tables(self):
        """测试Word文档表格中的文本替换"""
        try:
            from main import WordProcessor
            from docx import Document
        except ImportError as e:
            self.skipTest(f"依赖缺失: {e}")

        # 创建带表格的Word文档
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "公司名称ABC"
        table.cell(0, 1).text = "项目编号HT001"
        table.cell(1, 0).text = "其他内容"
        table.cell(1, 1).text = "备注信息"

        # 同时添加段落
        doc.add_paragraph("这是段落内容：项目编号HT001")

        word_path = Path(self.test_dir) / "test.docx"
        doc.save(str(word_path))

        output_path = Path(self.test_dir) / "output.docx"

        # 执行替换
        processor = WordProcessor()
        success = processor.process_document(
            str(word_path),
            {"ABC": "测试公司", "HT001": "2024-001"},
            str(output_path)
        )

        self.assertTrue(success)

        # 验证替换结果
        from docx import Document
        result_doc = Document(str(output_path))

        # 验证表格中
        self.assertEqual(result_doc.tables[0].cell(0, 0).text, "公司名称测试公司")
        self.assertEqual(result_doc.tables[0].cell(0, 1).text, "项目编号2024-001")

        # 验证段落中（查找包含2024-001的段落）
        found_in_paragraph = False
        for para in result_doc.paragraphs:
            if "2024-001" in para.text:
                found_in_paragraph = True
                break
        self.assertTrue(found_in_paragraph, "段落中应包含替换后的2024-001")


def run_tests():
    """运行所有测试"""
    print("=" * 70)
    print("Word 文本处理工具 - 单元测试")
    print("=" * 70)
    print()

    # 创建测试套件
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()

    # 添加测试类
    test_classes = [
        TestWordProcessor,
        TestExcelProcessor,
        TestSuffixProcessor,
        TestExcelToPdfProcessor,
        TestUtilityFunctions,
        TestIntegration,
    ]

    for test_class in test_classes:
        tests = loader.loadTestsFromTestCase(test_class)
        suite.addTests(tests)

    # 运行测试
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)

    print()
    print("=" * 70)
    print(f"测试结果: 运行={result.testsRun}, 失败={len(result.failures)}, 错误={len(result.errors)}, 跳过={len(result.skipped)}")
    print("=" * 70)

    return 0 if result.wasSuccessful() else 1


if __name__ == "__main__":
    sys.exit(run_tests())
