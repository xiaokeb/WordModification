#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word 文本处理工具
功能：批量替换 Word 文档中的文本，保持原有格式不变；记录编号并写入 Excel 文档
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import threading
import os
import json


class WordProcessor:
    """Word 文档处理类"""

    MAX_FILE_SIZE = 50 * 1024 * 1024
    MAX_REPLACE_TEXT_LENGTH = 10000

    def __init__(self):
        pass

    def get_word_files(self, path):
        """获取指定路径下的所有 Word 文件"""
        word_files = []
        try:
            if not path:
                return []
            
            if isinstance(path, (str, Path)):
                path = Path(path)
                if not path.exists():
                    return []
                
                try:
                    path = path.resolve()
                except Exception:
                    return []
                
                if path.is_dir():
                    for ext in ['*.doc', '*.docx']:
                        for file in path.glob(ext):
                            try:
                                if self._validate_file(file):
                                    word_files.append(str(file))
                            except (OSError, IOError):
                                continue
                elif path.is_file() and path.suffix.lower() in ['.doc', '.docx']:
                    try:
                        if self._validate_file(path):
                            word_files.append(str(path))
                    except (OSError, IOError):
                        pass
            elif isinstance(path, list):
                for p in path:
                    p = Path(p)
                    if p.is_file() and p.suffix.lower() in ['.doc', '.docx']:
                        try:
                            if self._validate_file(p):
                                word_files.append(str(p))
                        except (OSError, IOError):
                            continue
        except Exception:
            return []

        return word_files

    def _validate_file(self, file_path):
        """验证文件是否合法（存在、可访问、大小合规）"""
        if not file_path.exists():
            return False
        
        if not file_path.is_file():
            return False
        
        if file_path.stat().st_size > self.MAX_FILE_SIZE:
            return False
        
        if not os.access(file_path, os.R_OK):
            return False
        
        return True

    def _normalize_path(self, path_str):
        """规范化路径，防止路径遍历攻击"""
        try:
            path_obj = Path(path_str)
            if not path_obj.exists():
                return None
            
            resolved_path = path_obj.resolve()
            
            if '..' in str(path_obj):
                return None
            
            return resolved_path
        except Exception:
            return None

    def replace_text_in_paragraph(self, paragraph, replacements):
        """在段落中替换文本，保持格式不变"""
        sorted_replacements = sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True)
        
        for old_text, new_text in sorted_replacements:
            if not old_text:
                continue
            
            if len(old_text) > self.MAX_REPLACE_TEXT_LENGTH:
                continue
            if new_text and len(new_text) > self.MAX_REPLACE_TEXT_LENGTH:
                continue

            paragraph_text = paragraph.text
            if old_text not in paragraph_text:
                continue

            self._replace_text_preserve_format(paragraph, old_text, new_text)

    def _replace_text_preserve_format(self, paragraph, old_text, new_text):
        """替换文本并保持格式的核心方法"""
        runs = paragraph.runs
        if not runs:
            return

        full_text = paragraph.text
        start_pos = full_text.find(old_text)

        if start_pos == -1:
            return

        end_pos = start_pos + len(old_text)

        run_positions = self._calculate_run_positions(runs)
        affected_runs = self._find_affected_runs(run_positions, start_pos, end_pos)

        if not affected_runs:
            return

        if len(affected_runs) == 1:
            self._replace_single_run(affected_runs[0], start_pos, end_pos, new_text)
        else:
            self._replace_multi_runs(affected_runs, start_pos, end_pos, new_text)

    def _calculate_run_positions(self, runs):
        """计算每个 run 的位置范围"""
        run_positions = []
        current_pos = 0

        for i, run in enumerate(runs):
            run_len = len(run.text)
            run_positions.append((i, current_pos, current_pos + run_len, run))
            current_pos += run_len

        return run_positions

    def _find_affected_runs(self, run_positions, start_pos, end_pos):
        """查找受替换影响的 runs"""
        affected_runs = []
        for i, run_start, run_end, run in run_positions:
            if run_start < end_pos and run_end > start_pos:
                affected_runs.append((i, run_start, run_end, run))
        return affected_runs

    def _replace_single_run(self, run_info, start_pos, end_pos, new_text):
        """替换单个 run 中的文本"""
        idx, run_start, run_end, run = run_info
        local_start = start_pos - run_start
        local_end = end_pos - run_start
        run.text = run.text[:local_start] + new_text + run.text[local_end:]

    def _replace_multi_runs(self, affected_runs, start_pos, end_pos, new_text):
        """替换跨越多个 runs 的文本"""
        first_idx, first_start, first_end, first_run = affected_runs[0]
        last_idx, last_start, last_end, last_run = affected_runs[-1]

        first_local_start = start_pos - first_start
        last_local_end = end_pos - last_start

        first_run.text = first_run.text[:first_local_start] + new_text + last_run.text[last_local_end:]

        runs_to_remove = []
        for i in range(1, len(affected_runs)):
            idx, run_start, run_end, run = affected_runs[i]
            runs_to_remove.append(run)

        for run in runs_to_remove:
            try:
                run._element.getparent().remove(run._element)
            except Exception:
                run.text = ""

    def _copy_run_format(self, source_run, target_run):
        """复制 run 的格式"""
        if source_run.bold:
            target_run.bold = True
        if source_run.italic:
            target_run.italic = True
        if source_run.underline:
            target_run.underline = source_run.underline
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def process_document(self, doc_path, replacements, output_path, callback=None):
        """处理单个 Word 文档"""
        try:
            if not doc_path.lower().endswith('.docx'):
                return False
            
            file_size = Path(doc_path).stat().st_size
            if file_size > self.MAX_FILE_SIZE:
                print(f"文件过大：{doc_path} ({file_size / 1024 / 1024:.2f}MB)")
                return False
            
            output_path_obj = Path(output_path).resolve()
            
            if not str(output_path_obj).lower().endswith('.docx'):
                return False

            doc = Document(doc_path)

            for paragraph in doc.paragraphs:
                self.replace_text_in_paragraph(paragraph, replacements)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, replacements)

            doc.save(str(output_path_obj))

            if callback:
                callback()

            return True
        except FileNotFoundError as e:
            print(f"文件不存在 {doc_path}: {str(e)}")
            return False
        except PermissionError as e:
            print(f"权限错误 {doc_path}: {str(e)}")
            return False
        except Exception as e:
            print(f"处理文档失败 {doc_path}: {str(e)}")
            import traceback
            traceback.print_exc()
            return False


class ExcelProcessor:
    """Excel 文档处理类"""

    def __init__(self):
        pass

    def write_number_to_excel(self, excel_path, number_text, use_fixed_mode=False, fixed_number="", number_format="", callback=None):
        """
        将编号写入 Excel 文件
        
        Args:
            excel_path: Excel 文件路径
            number_text: 生成的编号文本
            use_fixed_mode: 是否使用固定编号模式
            fixed_number: 固定编号内容
            number_format: 编号格式
            callback: 回调函数
        """
        try:
            from openpyxl import load_workbook
            from openpyxl.styles import Font, Alignment

            wb = load_workbook(excel_path)
            ws = wb.active

            ws.insert_rows(1)

            ws.merge_cells('E1:F1')
            cell_e1 = ws['E1']
            cell_e1.font = Font(name='Times New Roman', bold=True)
            cell_e1.alignment = Alignment(horizontal='center', vertical='center')

            if use_fixed_mode and fixed_number:
                cell_e1.value = fixed_number

                ws.merge_cells('E2:F2')
                cell_e2 = ws['E2']
                cell_e2.value = number_format
                cell_e2.font = Font(name='Times New Roman', bold=True)
                cell_e2.alignment = Alignment(horizontal='center', vertical='center')
            else:
                cell_e1.value = number_text

            wb.save(excel_path)

            if callback:
                callback()

            return True
        except ImportError:
            print("缺少 openpyxl 库，请安装：pip install openpyxl")
            return False
        except Exception as e:
            print(f"处理 Excel 文件失败 {excel_path}: {str(e)}")
            import traceback
            traceback.print_exc()
            return False


class WordProcessorGUI:
    """Word 处理工具 GUI 类"""

    def __init__(self, root):
        """初始化 GUI 界面"""
        self.root = root
        self.root.title("Word 文本处理工具")
        
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        
        window_width = min(1000, screen_width - 100)
        window_height = min(800, screen_height - 100)
        
        self.root.geometry(f"{window_width}x{window_height}")
        
        self.processor = WordProcessor()
        self.excel_processor = ExcelProcessor()
        self.selected_files = []
        self.output_folder = ""
        
        self.settings_file = Path(__file__).parent / "settings.json"
        self.number_settings = self._load_number_settings()

        self._create_tabbed_interface()

    def _load_number_settings(self):
        """从本地文件加载编号设置"""
        try:
            if self.settings_file.exists():
                with open(self.settings_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception:
            pass
        
        return {
            "prefix": "HT合同编号",
            "current_number": 1,
            "number_format": "HT合同编号-年份-XM001-01",
            "fixed_number": ""
        }

    def _save_number_settings(self):
        """保存编号设置到本地文件"""
        try:
            with open(self.settings_file, 'w', encoding='utf-8') as f:
                json.dump(self.number_settings, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存设置失败: {e}")

    def _generate_number(self):
        """生成编号，序号自动递增，超过99后自动变为3位数"""
        number = self.number_settings["current_number"]
        if number <= 99:
            number_str = str(number).zfill(2)
        else:
            number_str = str(number)
        
        number_format = self.number_settings["number_format"]
        full_number = number_format.replace("01", number_str)
        
        self.number_settings["current_number"] = number + 1
        self._save_number_settings()
        
        return full_number

    def _create_tabbed_interface(self):
        """创建选项卡容器界面"""
        main_frame = ttk.Frame(self.root, padding="5")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.tab_text_replace = ttk.Frame(self.notebook, padding="5")
        self.tab_number_record = ttk.Frame(self.notebook, padding="5")
        
        self.notebook.add(self.tab_text_replace, text="文本替换")
        self.notebook.add(self.tab_number_record, text="记录编号")
        
        self._create_text_replace_tab()
        self._create_number_record_tab()

    # ==================== 文本替换功能 ====================

    def _create_text_replace_tab(self):
        """创建文本替换选项卡的所有组件"""
        self._create_file_section(self.tab_text_replace)
        self._create_file_list(self.tab_text_replace)
        self._create_file_buttons(self.tab_text_replace)
        self._create_replacement_section(self.tab_text_replace)
        self._create_output_section(self.tab_text_replace)
        self._create_progress_section(self.tab_text_replace)
        self._create_button_section(self.tab_text_replace)

    def _create_file_section(self, parent):
        """创建文件选择区域：包含浏览文件和浏览文件夹按钮"""
        file_frame = ttk.LabelFrame(parent, text="文件选择（可选择文件夹或文件）", padding="5")
        file_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        file_frame.columnconfigure(0, weight=1)

        browse_button_frame = ttk.Frame(file_frame)
        browse_button_frame.grid(row=0, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(0, 5))
        browse_button_frame.columnconfigure(1, weight=1)

        ttk.Label(browse_button_frame, text="选择文件/文件夹:").grid(row=0, column=0, sticky=tk.W)

        self.file_path_var = tk.StringVar()
        file_path_entry = ttk.Entry(browse_button_frame, textvariable=self.file_path_var, state='readonly')
        file_path_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(5, 5))

        browse_btn = ttk.Button(browse_button_frame, text="浏览文件", command=self._browse_files)
        browse_btn.grid(row=0, column=2, padx=(0, 5))

        browse_folder_btn = ttk.Button(browse_button_frame, text="浏览文件夹", command=self._browse_folder)
        browse_folder_btn.grid(row=0, column=3)

        self.file_count_label = ttk.Label(file_frame, text="未选择文件", foreground="gray")
        self.file_count_label.grid(row=1, column=0, sticky=tk.W, pady=(0, 5))

        self.file_path_detail_label = ttk.Label(
            file_frame, 
            text="", 
            foreground="gray", 
            font=("Microsoft YaHei UI", 9), 
            anchor='w'
        )
        self.file_path_detail_label.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 5))

    def _create_file_list(self, parent):
        """创建已选文件列表：显示已选择的 Word 文件"""
        file_list_frame = ttk.LabelFrame(parent, text="已选文件列表", padding="5")
        file_list_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        file_list_frame.columnconfigure(0, weight=1)
        file_list_frame.rowconfigure(0, weight=1)

        file_columns = ("filename",)
        self.file_tree = ttk.Treeview(file_list_frame, columns=file_columns, show="headings", height=4)
        self.file_tree.heading("filename", text="文件名")
        self.file_tree.column("filename", width=800, minwidth=800)
        
        style = ttk.Style()
        style.configure("FileTreeview", rowheight=30, font=("Microsoft YaHei UI", 9))
        style.configure("FileTreeview.Heading", font=("Microsoft YaHei UI", 9, "bold"))
        
        file_scrollbar = ttk.Scrollbar(file_list_frame, orient=tk.VERTICAL, command=self.file_tree.yview)
        self.file_tree.configure(yscrollcommand=file_scrollbar.set)
        
        self.file_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        file_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.file_tree.bind('<Double-1>', lambda e: self._remove_selected_file())

    def _create_file_buttons(self, parent):
        """创建文件操作按钮：添加和删除文件"""
        file_button_frame = ttk.Frame(parent)
        file_button_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        
        add_file_btn = ttk.Button(
            file_button_frame,
            text="+",
            width=65,
            command=self._add_file
        )
        add_file_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        remove_file_btn = ttk.Button(
            file_button_frame,
            text="-",
            width=65,
            command=self._remove_selected_file
        )
        remove_file_btn.pack(side=tk.LEFT)

    def _create_replacement_section(self, parent):
        """创建替换规则设置区域：显示可编辑的替换规则表格"""
        replacement_frame = ttk.LabelFrame(parent, text="替换规则设置", padding="5")
        replacement_frame.grid(row=3, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 5))
        replacement_frame.columnconfigure(0, weight=1)
        replacement_frame.rowconfigure(0, weight=1)

        self.replacement_entries = {}
        
        tree_frame = ttk.Frame(replacement_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        columns = ("category", "original", "replacement")
        tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=4)
        
        tree.heading("category", text="类别")
        tree.heading("original", text="查找")
        tree.heading("replacement", text="替换")
        
        tree.column("category", width=120, anchor='center', minwidth=120)
        tree.column("original", width=350, minwidth=350)
        tree.column("replacement", width=350, minwidth=350)
        
        style = ttk.Style()
        style.configure("Treeview", rowheight=30, font=("Microsoft YaHei UI", 9))
        style.configure("Treeview.Heading", font=("Microsoft YaHei UI", 9, "bold"))
        
        tree.tag_configure("category_style", foreground="#808080")
        tree.tag_configure("editable_style", foreground="#000000")

        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)

        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        tree.bind('<Double-1>', lambda e: self._edit_cell(tree))

        button_frame = ttk.Frame(replacement_frame)
        button_frame.pack(fill=tk.X, pady=(5, 0))

        reset_btn = ttk.Button(
            button_frame,
            text="重置为默认值",
            width=42,
            command=lambda: self._reset_to_defaults(tree)
        )
        reset_btn.pack(side=tk.LEFT)

        self.replacement_entries["tree"] = tree
        
        self._init_default_rules(tree)

    def _create_output_section(self, parent):
        """创建保存设置区域：选择输出文件保存路径"""
        output_frame = ttk.LabelFrame(parent, text="保存设置", padding="5")
        output_frame.grid(row=4, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        output_frame.columnconfigure(1, weight=1)

        ttk.Label(output_frame, text="保存路径:").grid(row=0, column=0, sticky=tk.W)

        self.output_path_var = tk.StringVar()
        output_path_entry = ttk.Entry(output_frame, textvariable=self.output_path_var, state='readonly')
        output_path_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(5, 5))

        output_browse_btn = ttk.Button(output_frame, text="浏览", command=self._browse_output)
        output_browse_btn.grid(row=0, column=2)

        self.output_info_label = ttk.Label(
            output_frame,
            text="未选择保存路径时将按规则自动创建文件夹，默认在同级目录下新建修改版文件夹",
            foreground="gray",
            font=("Microsoft YaHei UI", 9)
        )
        self.output_info_label.grid(row=1, column=0, columnspan=3, sticky=tk.W, pady=(5, 0))

    def _create_progress_section(self, parent):
        """创建处理进度区域：显示处理进度条和状态信息"""
        progress_frame = ttk.LabelFrame(parent, text="处理进度", padding="5")
        progress_frame.grid(row=5, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        progress_frame.columnconfigure(0, weight=1)

        self.progress_var = tk.IntVar()
        progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100)
        progress_bar.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 5))

        self.progress_label = ttk.Label(progress_frame, text="就绪", foreground="gray")
        self.progress_label.grid(row=1, column=0, sticky=tk.W)

    def _create_button_section(self, parent):
        """创建操作按钮区域：开始处理和清空设置按钮"""
        button_frame = ttk.Frame(parent)
        button_frame.grid(row=6, column=0, pady=(5, 0))

        self.start_btn = ttk.Button(
            button_frame,
            text="开始处理",
            command=self._start_processing,
            width=20
        )
        self.start_btn.pack(side=tk.LEFT, padx=(0, 10))

        clear_btn = ttk.Button(
            button_frame,
            text="清空设置",
            command=self._clear_all,
            width=20
        )
        clear_btn.pack(side=tk.LEFT)

    def _browse_files(self):
        """浏览并选择多个 Word 文件"""
        files = filedialog.askopenfilenames(
            title="选择 Word 文件",
            filetypes=[("Word 文件", "*.docx *.doc"), ("所有文件", "*.*")]
        )
        if files:
            self.selected_files = list(files)
            self.file_path_var.set(f"已选择 {len(self.selected_files)} 个文件")
            self.file_count_label.config(
                text=f"已选择 {len(self.selected_files)} 个文件",
                foreground="blue"
            )
            
            paths = "\n".join(self.selected_files[:5])
            if len(self.selected_files) > 5:
                paths += f"\n... 及其他 {len(self.selected_files) - 5} 个文件"
            self.file_path_detail_label.config(text=paths)
            
            self._update_file_list()

    def _browse_folder(self):
        """浏览并选择包含 Word 文件的文件夹"""
        folder = filedialog.askdirectory(title="选择包含 Word 文件的文件夹")
        if folder:
            word_files = self.processor.get_word_files(folder)
            if word_files:
                self.selected_files = word_files
                self.file_path_var.set(f"已选择文件夹：{folder}，共 {len(self.selected_files)} 个文件")
                self.file_count_label.config(
                    text=f"已选择 {len(self.selected_files)} 个文件",
                    foreground="blue"
                )
                
                paths = "\n".join(self.selected_files[:5])
                if len(self.selected_files) > 5:
                    paths += f"\n... 及其他 {len(self.selected_files) - 5} 个文件"
                self.file_path_detail_label.config(text=paths)
                
                self._update_file_list()
            else:
                messagebox.showwarning("警告", "该文件夹下没有找到 Word 文件")

    def _add_file(self):
        """添加额外的 Word 文件到已选列表"""
        files = filedialog.askopenfilenames(
            title="添加 Word 文件",
            filetypes=[("Word 文件", "*.docx *.doc"), ("所有文件", "*.*")]
        )
        if files:
            for file in files:
                if file not in self.selected_files:
                    self.selected_files.append(file)
            
            self.file_path_var.set(f"已选择 {len(self.selected_files)} 个文件")
            self.file_count_label.config(
                text=f"已选择 {len(self.selected_files)} 个文件",
                foreground="blue"
            )
            
            paths = "\n".join(self.selected_files[:5])
            if len(self.selected_files) > 5:
                paths += f"\n... 及其他 {len(self.selected_files) - 5} 个文件"
            self.file_path_detail_label.config(text=paths)
            
            self._update_file_list()

    def _remove_selected_file(self):
        """从已选列表中删除选中的文件"""
        selected = self.file_tree.selection()
        if selected:
            for item in selected:
                values = self.file_tree.item(item, 'values')
                if values:
                    file_path = values[0]
                    if file_path in self.selected_files:
                        self.selected_files.remove(file_path)
                    self.file_tree.delete(item)
            
            if self.selected_files:
                self.file_count_label.config(
                    text=f"已选择 {len(self.selected_files)} 个文件",
                    foreground="blue"
                )
                self.file_path_var.set(f"已选择 {len(self.selected_files)} 个文件")
            else:
                self.file_count_label.config(
                    text="未选择文件",
                    foreground="gray"
                )
                self.file_path_var.set("")
                self.file_path_detail_label.config(text="")

    def _update_file_list(self):
        """更新文件列表显示"""
        self.file_tree.delete(*self.file_tree.get_children())
        for file_path in self.selected_files:
            file_name = Path(file_path).name
            self.file_tree.insert("", tk.END, values=(file_path,))

    def _browse_output(self):
        """浏览并选择输出文件保存路径"""
        folder = filedialog.askdirectory(title="选择保存路径")
        if folder:
            try:
                self.output_folder = folder
                self.output_path_var.set(folder)
                self.output_info_label.config(
                    text="已选择自定义保存路径",
                    foreground="green"
                )
            except Exception as e:
                messagebox.showerror("错误", f"无法设置保存路径：{str(e)}")

    def _init_default_rules(self, tree):
        """初始化默认替换规则"""
        self.default_rules = [
            ("项目编号", "HT合同编号-年份-XM001-01", "HT合同编号-年份-XM001-01"),
            ("公司名称", "", ""),
            ("系统名称", "", ""),
            ("人员名称", "", "")
        ]
        
        for category, original, replacement in self.default_rules:
            tree.insert("", tk.END, values=(category, original, replacement), tags=('category_style', 'editable_style'))
    
    def _reset_to_defaults(self, tree):
        """重置替换规则为默认值"""
        tree.delete(*tree.get_children())
        self._init_default_rules(tree)
    
    def _edit_cell(self, tree):
        """编辑表格单元格（双击触发）"""
        selected = tree.selection()
        if not selected:
            return
        
        item = selected[0]
        
        x = tree.winfo_pointerx() - tree.winfo_rootx()
        y = tree.winfo_pointery() - tree.winfo_rooty()
        
        row = tree.identify_row(y)
        col = tree.identify_column(x)
        
        if not row or not col:
            return
        
        col_index = int(col.replace('#', '')) - 1
        
        if col_index == 0:
            return
        
        values = list(tree.item(item, 'values'))
        
        bbox = tree.bbox(item, col)
        if bbox is None:
            return
        
        cell_x, cell_y, cell_width, cell_height = bbox
        
        entry = ttk.Entry(tree, font=("Microsoft YaHei UI", 9))
        entry.place(x=cell_x, y=cell_y, width=cell_width, height=cell_height)
        
        current_value = values[col_index] if col_index < len(values) else ""
        entry.insert(0, current_value)
        entry.focus()
        
        self._bind_entry_events(entry, tree, item, col_index, values)
    
    def _bind_entry_events(self, entry, tree, item, col_index, values):
        """绑定编辑框事件（回车保存、失焦保存、ESC取消）"""
        def save_edit(event=None):
            new_value = entry.get()
            while len(values) <= col_index:
                values.append("")
            values[col_index] = new_value
            tree.item(item, values=tuple(values))
            entry.destroy()
        
        entry.bind('<Return>', save_edit)
        entry.bind('<FocusOut>', save_edit)
        entry.bind('<Escape>', lambda e: entry.destroy())
    
    def _get_all_replacements(self):
        """获取所有替换规则并转换为字典"""
        replacements = {}
        tree = self.replacement_entries["tree"]
        
        for item in tree.get_children():
            values = tree.item(item, 'values')
            if len(values) >= 3:
                category, original, replacement = values[0], values[1], values[2]
                if original and original.strip():
                    replacements[original.strip()] = replacement.strip() if replacement else ""
        
        return replacements

    def _clear_all(self):
        """清空所有设置并恢复默认状态"""
        self.selected_files = []
        self.output_folder = ""
        self.file_path_var.set("")
        self.output_path_var.set("")
        self.file_count_label.config(
            text="未选择文件",
            foreground="gray"
        )
        self.output_info_label.config(
            text="未选择保存路径时将按规则自动创建文件夹，默认在同级目录下新建修改版文件夹",
            foreground="gray"
        )
        self.progress_var.set(0)
        self.progress_label.config(text="就绪")

        self.file_tree.delete(*self.file_tree.get_children())

        tree = self.replacement_entries["tree"]
        self._reset_to_defaults(tree)

    def _start_processing(self):
        """开始处理 Word 文档（在后台线程中执行）"""
        if not self.selected_files:
            messagebox.showwarning("警告", "请先选择要处理的 Word 文件")
            return

        replacements = self._get_all_replacements()
        if not replacements:
            messagebox.showwarning("警告", "请至少添加一条替换规则")
            return

        self.start_btn.config(state='disabled')

        def process_thread():
            total_files = len(self.selected_files)
            success_count = 0

            for i, file_path in enumerate(self.selected_files):
                file_name = Path(file_path).name
                base_name = Path(file_path).stem

                self.root.after(0, lambda f=file_name, idx=i + 1, total=total_files:
                    self.progress_label.config(
                        text=f"正在处理：{f} ({idx}/{total})"
                    )
                )

                if self.output_folder:
                    output_dir = Path(self.output_folder)
                else:
                    file_dir = Path(file_path).parent
                    output_dir = file_dir / "修改版"

                output_dir.mkdir(parents=True, exist_ok=True)

                output_path = output_dir / f"{base_name}-修改版.docx"

                success = self.processor.process_document(
                    file_path,
                    replacements,
                    str(output_path)
                )

                if success:
                    success_count += 1

                progress = int(((i + 1) / total_files) * 100)
                self.root.after(0, lambda p=progress: self.progress_var.set(p))

            self.root.after(0, lambda: self._processing_complete(success_count, total_files))

        thread = threading.Thread(target=process_thread)
        thread.daemon = True
        thread.start()

    def _processing_complete(self, success_count, total_files):
        """处理完成后的回调函数"""
        self.start_btn.config(state='normal')

        message = f"处理完成！\n成功：{success_count}/{total_files}"
        if success_count == total_files:
            message += "\n\n修改完成，已保存"
            messagebox.showinfo("完成", message)
        else:
            messagebox.showwarning("完成", message)

        self.progress_label.config(text="处理完成")

    # ==================== 记录编号功能 ====================

    def _create_number_record_tab(self):
        """创建记录编号选项卡的所有组件"""
        self._create_number_mode_section(self.tab_number_record)
        self._create_number_format_section(self.tab_number_record)
        self._create_number_fixed_section(self.tab_number_record)
        self._create_number_file_section(self.tab_number_record)
        self._create_number_progress_section(self.tab_number_record)
        self._create_number_button_section(self.tab_number_record)

    def _create_number_mode_section(self, parent):
        """创建编号模式选择区域：自动编号模式和固定编号模式"""
        mode_frame = ttk.LabelFrame(parent, text="编号模式", padding="5")
        mode_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        mode_frame.columnconfigure(0, weight=1)

        self.number_mode_var = tk.StringVar(value="auto")
        
        auto_mode_radio = ttk.Radiobutton(
            mode_frame, 
            text="自动编号模式（仅插入编号格式行）", 
            variable=self.number_mode_var, 
            value="auto",
            command=self._on_mode_change
        )
        auto_mode_radio.grid(row=0, column=0, sticky=tk.W, padx=(0, 20))

        fixed_mode_radio = ttk.Radiobutton(
            mode_frame, 
            text="固定编号模式（插入固定编号+编号格式行）", 
            variable=self.number_mode_var, 
            value="fixed",
            command=self._on_mode_change
        )
        fixed_mode_radio.grid(row=1, column=0, sticky=tk.W)

    def _create_number_format_section(self, parent):
        """创建编号格式设置区域：可编辑的编号格式输入框和预览"""
        format_frame = ttk.LabelFrame(parent, text="编号格式设置", padding="5")
        format_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        format_frame.columnconfigure(1, weight=1)

        ttk.Label(format_frame, text="编号格式:").grid(row=0, column=0, sticky=tk.W, padx=(0, 5))

        self.number_format_var = tk.StringVar(value=self.number_settings.get("number_format", "HT合同编号-年份-XM001-01"))
        format_entry = ttk.Entry(format_frame, textvariable=self.number_format_var)
        format_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(0, 5))

        ttk.Label(format_frame, text="说明：使用 01 作为序号占位符，序号自动递增，超过99后自动变为3位数", 
                 foreground="gray", font=("Microsoft YaHei UI", 9)).grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=(5, 0))

        preview_frame = ttk.Frame(format_frame)
        preview_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(5, 0))
        
        ttk.Label(preview_frame, text="当前编号预览:").pack(side=tk.LEFT, padx=(0, 5))
        self.current_number_var = tk.StringVar()
        self.current_number_var.set(self._generate_preview())
        preview_label = ttk.Entry(preview_frame, textvariable=self.current_number_var, state='readonly')
        preview_label.pack(side=tk.LEFT, fill=tk.X, expand=True)

    def _create_number_fixed_section(self, parent):
        """创建固定编号输入区域：仅在固定编号模式下使用"""
        fixed_frame = ttk.LabelFrame(parent, text="固定编号（仅在固定编号模式下使用）", padding="5")
        fixed_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        fixed_frame.columnconfigure(1, weight=1)

        ttk.Label(fixed_frame, text="固定编号:").grid(row=0, column=0, sticky=tk.W, padx=(0, 5))

        self.fixed_number_var = tk.StringVar(value=self.number_settings.get("fixed_number", ""))
        fixed_entry = ttk.Entry(fixed_frame, textvariable=self.fixed_number_var)
        fixed_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(0, 5))

        self.fixed_info_label = ttk.Label(
            fixed_frame,
            text="固定编号将写入 E1:F1 合并单元格，编号格式将写入 E2:F2 合并单元格",
            foreground="gray",
            font=("Microsoft YaHei UI", 9)
        )
        self.fixed_info_label.grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=(5, 0))

    def _on_mode_change(self):
        """编号模式切换时的处理函数"""
        if self.number_mode_var.get() == "fixed":
            self.fixed_number_var.set(self.number_settings.get("fixed_number", ""))
        else:
            self.fixed_number_var.set("")

    def _create_number_file_section(self, parent):
        """创建 Excel 文件选择区域：选择要写入编号的 Excel 文件"""
        file_frame = ttk.LabelFrame(parent, text="选择 Excel 文件", padding="5")
        file_frame.grid(row=3, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        file_frame.columnconfigure(0, weight=1)

        browse_button_frame = ttk.Frame(file_frame)
        browse_button_frame.grid(row=0, column=0, sticky=(tk.W, tk.E))
        browse_button_frame.columnconfigure(1, weight=1)

        ttk.Label(browse_button_frame, text="Excel 文件:").grid(row=0, column=0, sticky=tk.W)

        self.excel_file_path_var = tk.StringVar()
        excel_file_entry = ttk.Entry(browse_button_frame, textvariable=self.excel_file_path_var, state='readonly')
        excel_file_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(5, 5))

        browse_excel_btn = ttk.Button(browse_button_frame, text="浏览", command=self._browse_excel_file)
        browse_excel_btn.grid(row=0, column=2)

        self.excel_file_info_label = ttk.Label(
            file_frame,
            text="未选择文件",
            foreground="gray"
        )
        self.excel_file_info_label.grid(row=1, column=0, sticky=tk.W, pady=(5, 0))

    def _create_number_progress_section(self, parent):
        """创建记录编号处理进度区域：显示写入进度和状态"""
        progress_frame = ttk.LabelFrame(parent, text="处理进度", padding="5")
        progress_frame.grid(row=4, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        progress_frame.columnconfigure(0, weight=1)

        self.number_progress_var = tk.IntVar()
        number_progress_bar = ttk.Progressbar(progress_frame, variable=self.number_progress_var, maximum=100)
        number_progress_bar.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 5))

        self.number_progress_label = ttk.Label(progress_frame, text="就绪", foreground="gray")
        self.number_progress_label.grid(row=1, column=0, sticky=tk.W)

    def _create_number_button_section(self, parent):
        """创建记录编号操作按钮区域：写入编号和重置编号按钮"""
        button_frame = ttk.Frame(parent)
        button_frame.grid(row=5, column=0, pady=(5, 0))

        self.write_number_btn = ttk.Button(
            button_frame,
            text="写入编号",
            command=self._start_write_number,
            width=20
        )
        self.write_number_btn.pack(side=tk.LEFT, padx=(0, 10))

        reset_number_btn = ttk.Button(
            button_frame,
            text="重置编号",
            command=self._reset_number,
            width=20
        )
        reset_number_btn.pack(side=tk.LEFT)

    def _generate_preview(self):
        """生成当前编号预览（不递增序号）"""
        number = self.number_settings["current_number"]
        if number <= 99:
            number_str = str(number).zfill(2)
        else:
            number_str = str(number)
        
        number_format = self.number_settings.get("number_format", "HT合同编号-年份-XM001-01")
        return number_format.replace("01", number_str)

    def _browse_excel_file(self):
        """浏览并选择 Excel 文件"""
        file = filedialog.askopenfilename(
            title="选择 Excel 文件",
            filetypes=[("Excel 文件", "*.xlsx *.xls"), ("所有文件", "*.*")]
        )
        if file:
            self.excel_file_path_var.set(file)
            self.excel_file_info_label.config(
                text=f"已选择：{Path(file).name}",
                foreground="blue"
            )

    def _start_write_number(self):
        """开始写入编号到 Excel 文件（在后台线程中执行）"""
        excel_file = self.excel_file_path_var.get()
        if not excel_file:
            messagebox.showwarning("警告", "请先选择要写入的 Excel 文件")
            return

        number_format = self.number_format_var.get().strip()
        if not number_format:
            messagebox.showwarning("警告", "请输入编号格式")
            return

        self.number_settings["number_format"] = number_format
        
        use_fixed_mode = self.number_mode_var.get() == "fixed"
        fixed_number = ""
        
        if use_fixed_mode:
            fixed_number = self.fixed_number_var.get().strip()
            if not fixed_number:
                messagebox.showwarning("警告", "请输入固定编号")
                return
            self.number_settings["fixed_number"] = fixed_number
        
        self._save_number_settings()

        self.write_number_btn.config(state='disabled')
        self.number_progress_label.config(text="正在处理...")
        self.number_progress_var.set(0)

        def write_thread():
            number_text = self._generate_number()
            
            self.root.after(0, lambda: self.number_progress_label.config(text=f"正在写入编号：{number_text}"))
            
            success = self.excel_processor.write_number_to_excel(
                excel_file,
                number_text,
                use_fixed_mode=use_fixed_mode,
                fixed_number=fixed_number,
                number_format=number_format
            )
            
            self.root.after(0, lambda: self.number_progress_var.set(100))
            
            if success:
                self.root.after(0, lambda: self._write_number_complete(number_text, use_fixed_mode, fixed_number))
            else:
                self.root.after(0, lambda: self._write_number_failed())

        thread = threading.Thread(target=write_thread)
        thread.daemon = True
        thread.start()

    def _write_number_complete(self, number_text, use_fixed_mode, fixed_number):
        """编号写入成功后的回调函数"""
        self.write_number_btn.config(state='normal')
        self.current_number_var.set(self._generate_preview())
        
        if use_fixed_mode:
            self.number_progress_label.config(text=f"写入成功！固定编号：{fixed_number}，格式编号：{number_text}")
            messagebox.showinfo("完成", f"编号写入成功！\n\n固定编号（E1:F1）：{fixed_number}\n格式编号（E2:F2）：{number_text}")
        else:
            self.number_progress_label.config(text=f"写入成功！编号：{number_text}")
            messagebox.showinfo("完成", f"编号写入成功！\n\n编号（E1:F1）：{number_text}")

    def _write_number_failed(self):
        """编号写入失败后的回调函数"""
        self.write_number_btn.config(state='normal')
        self.number_progress_label.config(text="写入失败", foreground="red")
        messagebox.showerror("错误", "编号写入失败，请检查文件是否被占用或格式是否正确")

    def _reset_number(self):
        """重置编号序号为 01"""
        if messagebox.askyesno("确认", "确定要重置编号吗？当前编号将恢复为 01"):
            self.number_settings["current_number"] = 1
            self._save_number_settings()
            self.current_number_var.set(self._generate_preview())
            messagebox.showinfo("完成", "编号已重置为 01")


def main():
    """主函数：启动应用程序"""
    root = tk.Tk()
    root.title("Word 文本处理工具")
    root.geometry("900x700")

    try:
        from tkinter import font
        default_font = font.Font(family="Microsoft YaHei UI", size=9)
        root.option_add("*Font", default_font)
    except:
        pass

    app = WordProcessorGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
