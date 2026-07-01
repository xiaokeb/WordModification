#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word文档批量处理工具
功能：
1. 文本替换 - 批量替换Word文档中的文本（支持正文、表格、页眉页脚）
2. 记录编号 - 向Excel文件写入编号信息
3. 后缀修改 - 批量删除文件名中的"-修改版"后缀
"""

import sys
import os
import json
from pathlib import Path
from docx import Document

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QTabWidget, QLabel, QLineEdit,
    QPushButton, QTableWidget, QTableWidgetItem, QProgressBar, QFileDialog,
    QMessageBox, QHBoxLayout, QVBoxLayout, QHeaderView,
    QScrollArea, QSizePolicy, QGroupBox, QAbstractItemView, QStyleFactory
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont, QColor


# ============================================================================
# Word文档处理器
# ============================================================================
class WordProcessor:
    """Word文档处理核心类，负责文本替换逻辑"""
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

    def __init__(self):
        pass

    def get_word_files(self, path):
        """获取指定路径下的所有Word文件列表
        
        Args:
            path: 文件或文件夹路径，支持字符串、Path对象或列表
            
        Returns:
            list: Word文件路径列表
        """
        word_files = []
        try:
            if not path:
                return []

            if isinstance(path, (str, Path)):
                path = Path(path)
                if not path.exists():
                    return []

                try:
                    path = path.resolve()
                except Exception:
                    return []

                if path.is_dir():
                    for ext in ['*.doc', '*.docx']:
                        for file in path.glob(ext):
                            try:
                                if self._validate_file(file):
                                    word_files.append(str(file))
                            except (OSError, IOError):
                                continue
                elif path.is_file() and path.suffix.lower() in ['.doc', '.docx']:
                    try:
                        if self._validate_file(path):
                            word_files.append(str(path))
                    except (OSError, IOError):
                        pass
            elif isinstance(path, list):
                for p in path:
                    p = Path(p)
                    if p.is_file() and p.suffix.lower() in ['.doc', '.docx']:
                        try:
                            if self._validate_file(p):
                                word_files.append(str(p))
                        except (OSError, IOError):
                            continue
        except Exception:
            return []

        return word_files

    def _validate_file(self, file_path):
        """验证文件是否可处理（存在、可读、大小限制）"""
        if not file_path.exists():
            return False
        if not file_path.is_file():
            return False
        if file_path.stat().st_size > self.MAX_FILE_SIZE:
            return False
        if not os.access(file_path, os.R_OK):
            return False
        return True

    def replace_text_in_paragraph(self, paragraph, replacements):
        """替换段落中的文本
        
        按替换规则长度降序排列，优先替换较长的文本，避免部分匹配问题。
        
        Args:
            paragraph: docx段落对象
            replacements: 替换规则字典 {原文本: 新文本}
        """
        # 按原文本长度降序排序，优先替换长文本
        sorted_replacements = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)

        for old_text, new_text in sorted_replacements:
            if not old_text:
                continue
            self._replace_all_occurrences(paragraph, old_text, new_text)

    def _replace_all_occurrences(self, paragraph, old_text, new_text):
        """替换段落中所有匹配的文本
        
        处理跨Run的文本替换，确保完整替换所有匹配项。
        
        Args:
            paragraph: docx段落对象
            old_text: 要查找的原文本
            new_text: 替换后的新文本
        """
        max_iterations = 100  # 防止无限循环
        iteration = 0
        
        while iteration < max_iterations:
            iteration += 1
            if not self._replace_single_occurrence(paragraph, old_text, new_text):
                break

    def _replace_single_occurrence(self, paragraph, old_text, new_text):
        """替换段落中单个匹配项
        
        核心替换逻辑：
        1. 合并所有Run的文本
        2. 查找目标文本位置
        3. 定位文本跨越的Run范围
        4. 在第一个Run中执行替换，删除后续Run中的相关部分
        
        Args:
            paragraph: docx段落对象
            old_text: 要查找的原文本
            new_text: 替换后的新文本
            
        Returns:
            bool: 是否成功执行替换
        """
        runs = paragraph.runs
        if not runs:
            return False

        # 合并所有Run的文本
        full_text = ''.join(run.text for run in runs)
        
        # 查找目标文本
        start_pos = full_text.find(old_text)
        if start_pos == -1:
            return False
            
        end_pos = start_pos + len(old_text)

        # 计算每个Run的文本位置范围
        run_positions = []
        current_pos = 0
        for i, run in enumerate(runs):
            run_len = len(run.text)
            run_positions.append({
                'index': i,
                'start': current_pos,
                'end': current_pos + run_len,
                'run': run
            })
            current_pos += run_len

        # 找到文本跨越的Run范围
        first_run_info = None
        last_run_info = None
        first_local_start = 0
        last_local_end = 0

        for pos_info in run_positions:
            run_start = pos_info['start']
            run_end = pos_info['end']
            
            # 检查Run是否与目标文本有重叠
            overlap_start = max(run_start, start_pos)
            overlap_end = min(run_end, end_pos)
            
            if overlap_start < overlap_end:
                if first_run_info is None:
                    first_run_info = pos_info
                    first_local_start = overlap_start - run_start
                last_run_info = pos_info
                last_local_end = overlap_end - run_start

        if first_run_info is None or last_run_info is None:
            return False

        # 在第一个Run中执行替换
        first_run = first_run_info['run']
        original_text = first_run.text
        first_run.text = original_text[:first_local_start] + new_text + original_text[last_local_end:]

        # 删除后续Run中被替换文本覆盖的部分
        if first_run_info['index'] != last_run_info['index']:
            for pos_info in run_positions:
                if pos_info['index'] <= first_run_info['index']:
                    continue
                if pos_info['index'] > last_run_info['index']:
                    break
                    
                run = pos_info['run']
                run_start = pos_info['start']
                run_end = pos_info['end']
                
                if pos_info['index'] == last_run_info['index']:
                    # 最后一个Run，保留后面的文本
                    run.text = run.text[last_local_end:]
                else:
                    # 中间的Run，全部清空
                    run.text = ""

        return True

    def process_document(self, doc_path, replacements, output_path, callback=None):
        """处理Word文档，执行文本替换
        
        处理范围包括：正文段落、表格、页眉、页脚
        
        Args:
            doc_path: 输入文档路径
            replacements: 替换规则字典
            output_path: 输出文档路径
            callback: 处理完成回调函数
            
        Returns:
            bool: 处理是否成功
        """
        try:
            if not doc_path.lower().endswith('.docx'):
                return False

            file_size = Path(doc_path).stat().st_size
            if file_size > self.MAX_FILE_SIZE:
                print(f"文件过大：{doc_path} ({file_size / 1024 / 1024:.2f}MB)")
                return False

            output_path_obj = Path(output_path).resolve()
            if not str(output_path_obj).lower().endswith('.docx'):
                return False

            doc = Document(doc_path)

            # 替换正文段落
            for paragraph in doc.paragraphs:
                self.replace_text_in_paragraph(paragraph, replacements)

            # 替换表格内容（支持嵌套表格）
            self._process_tables(doc.tables, replacements)

            # 替换页眉
            self._process_headers(doc.sections, replacements)

            # 替换页脚
            self._process_footers(doc.sections, replacements)

            doc.save(str(output_path_obj))

            if callback:
                callback()

            return True
        except FileNotFoundError as e:
            print(f"文件不存在 {doc_path}: {str(e)}")
            return False
        except PermissionError as e:
            print(f"权限错误 {doc_path}: {str(e)}")
            return False
        except Exception as e:
            print(f"处理文档失败 {doc_path}: {str(e)}")
            import traceback
            traceback.print_exc()
            return False

    def _process_tables(self, tables, replacements):
        """处理文档中的所有表格
        
        支持嵌套表格，递归处理每个单元格中的段落。
        
        Args:
            tables: 表格集合
            replacements: 替换规则字典
        """
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_text_in_paragraph(paragraph, replacements)
                    # 处理单元格中的嵌套表格
                    if cell.tables:
                        self._process_tables(cell.tables, replacements)

    def _process_headers(self, sections, replacements):
        """处理文档所有节的页眉
        
        Args:
            sections: 文档节集合
            replacements: 替换规则字典
        """
        for section in sections:
            header = section.header
            if header and header.paragraphs:
                for paragraph in header.paragraphs:
                    self.replace_text_in_paragraph(paragraph, replacements)
            # 处理页眉中的表格
            if header and header.tables:
                self._process_tables(header.tables, replacements)

    def _process_footers(self, sections, replacements):
        """处理文档所有节的页脚
        
        Args:
            sections: 文档节集合
            replacements: 替换规则字典
        """
        for section in sections:
            footer = section.footer
            if footer and footer.paragraphs:
                for paragraph in footer.paragraphs:
                    self.replace_text_in_paragraph(paragraph, replacements)
            # 处理页脚中的表格
            if footer and footer.tables:
                self._process_tables(footer.tables, replacements)


# ============================================================================
# Excel处理器
# ============================================================================
class ExcelProcessor:
    """Excel文件处理类，负责编号写入"""

    def __init__(self):
        pass

    def write_number_to_excel(self, excel_path, output_path, number_text, 
                              use_fixed_mode=False, fixed_number="", 
                              number_format="", callback=None):
        """向Excel文件写入编号信息
        
        Args:
            excel_path: 输入Excel文件路径
            output_path: 输出Excel文件路径
            number_text: 格式编号文本
            use_fixed_mode: 是否使用固定编号模式
            fixed_number: 固定编号内容
            number_format: 编号格式模板
            callback: 完成回调函数
            
        Returns:
            bool: 写入是否成功
        """
        try:
            from openpyxl import load_workbook
            from openpyxl.styles import Font, Alignment

            wb = load_workbook(excel_path)

            header_font = Font(name='Times New Roman', bold=True, size=11)
            center_align = Alignment(horizontal='center', vertical='center')

            if use_fixed_mode and fixed_number:
                sheet_num = 1
                for ws in wb.worksheets:
                    ws.insert_rows(1)

                    ws.merge_cells('E1:F1')
                    cell_e1 = ws['E1']
                    cell_e1.font = header_font
                    cell_e1.alignment = center_align
                    cell_e1.value = fixed_number

                    ws.merge_cells('E2:F2')
                    cell_e2 = ws['E2']
                    cell_e2.font = header_font
                    cell_e2.alignment = center_align
                    sheet_num_str = str(sheet_num).zfill(2) if sheet_num <= 99 else str(sheet_num)
                    cell_e2.value = number_format.replace("yy", sheet_num_str)
                    sheet_num += 1
            else:
                sheet_num = 1
                for ws in wb.worksheets:
                    ws.merge_cells('E1:F1')
                    cell_e1 = ws['E1']
                    cell_e1.font = header_font
                    cell_e1.alignment = center_align
                    sheet_num_str = str(sheet_num).zfill(2) if sheet_num <= 99 else str(sheet_num)
                    cell_e1.value = number_format.replace("yy", sheet_num_str)
                    sheet_num += 1

            wb.save(output_path)

            if callback:
                callback()

            return True
        except ImportError:
            print("缺少 openpyxl 库，请安装：pip install openpyxl")
            return False
        except Exception as e:
            print(f"处理 Excel 文件失败 {excel_path}: {str(e)}")
            import traceback
            traceback.print_exc()
            return False


# ============================================================================
# Excel转PDF处理器
# ============================================================================
class ExcelToPdfProcessor:
    """Excel转PDF处理类，负责页面设置和PDF导出"""

    def __init__(self):
        pass

    def export_to_pdf(self, excel_path, output_path, callback=None):
        """将Excel文件导出为PDF
        
        设置所有工作表：
        - 缩放：将所有列缩放为一页（FitToPagesWide=1, FitToPagesTall=0）
        - 页边距：上下1.3cm，左右0.9cm
        
        Args:
            excel_path: 输入Excel文件路径
            output_path: 输出PDF文件路径
            callback: 完成回调函数
            
        Returns:
            bool: 导出是否成功
        """
        try:
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()

            excel = win32com.client.Dispatch("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False

            wb = excel.Workbooks.Open(str(Path(excel_path).resolve()))

            for ws in wb.Worksheets:
                # 设置页面缩放：将所有列缩放为一页，行数不限
                ws.PageSetup.FitToPagesWide = 1
                ws.PageSetup.FitToPagesTall = 0
                ws.PageSetup.Zoom = False  # 必须设为False才能使FitToPages生效

                # 设置页边距（单位：厘米）
                ws.PageSetup.TopMargin = excel.CentimetersToPoints(1.3)
                ws.PageSetup.BottomMargin = excel.CentimetersToPoints(1.3)
                ws.PageSetup.LeftMargin = excel.CentimetersToPoints(0.9)
                ws.PageSetup.RightMargin = excel.CentimetersToPoints(0.9)

            # 选中所有工作表
            wb.Worksheets.Select()

            # 导出为PDF
            wb.ExportAsFixedFormat(0, str(Path(output_path).resolve()))  # 0 = xlTypePDF

            wb.Close(False)
            excel.Quit()

            pythoncom.CoUninitialize()

            if callback:
                callback()

            return True
        except ImportError:
            print("缺少 pywin32 库，请安装：pip install pywin32")
            return False
        except Exception as e:
            print(f"导出PDF失败 {excel_path}: {str(e)}")
            import traceback
            traceback.print_exc()
            try:
                excel.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            return False


# ============================================================================
# 文件后缀处理器
# ============================================================================
class SuffixProcessor:
    """文件后缀处理类，负责批量修改文件名"""

    def __init__(self):
        pass

    def remove_suffix_from_files(self, files, suffix="-修改版", callback=None):
        """批量删除文件名中所有的指定后缀（包括中间和末尾）
        
        Args:
            files: 文件路径列表
            suffix: 要删除的后缀，默认为"-修改版"
            callback: 进度回调 (当前索引, 总数, 文件名)
            
        Returns:
            tuple: (成功数量, 总数)
        """
        total = len(files)
        success_count = 0

        for i, file_path in enumerate(files):
            try:
                path = Path(file_path)
                if not path.exists():
                    continue

                original_name = path.stem
                extension = path.suffix
                
                # 删除文件名中所有的后缀（包括中间和末尾）
                new_stem = original_name.replace(suffix, "")
                
                # 只有当名称实际发生变化时才重命名
                if new_stem != original_name and new_stem:
                    new_name = new_stem + extension
                    new_path = path.parent / new_name
                    
                    # 避免文件名冲突
                    if new_path.exists():
                        counter = 1
                        while new_path.exists():
                            new_name = f"{new_stem}_{counter}{extension}"
                            new_path = path.parent / new_name
                            counter += 1
                    
                    path.rename(new_path)
                    success_count += 1

                if callback:
                    callback(i, total, path.name)
            except Exception as e:
                print(f"处理文件失败 {file_path}: {str(e)}")

        return success_count, total


# ============================================================================
# 后台处理线程
# ============================================================================
class ProcessingThread(QThread):
    """Word文档批量处理线程"""
    progress_updated = pyqtSignal(int, str)
    processing_complete = pyqtSignal(int, int)

    def __init__(self, processor, files, replacements, output_folder):
        """初始化处理线程
        
        Args:
            processor: WordProcessor实例
            files: 待处理文件列表
            replacements: 替换规则字典
            output_folder: 输出文件夹路径
        """
        super().__init__()
        self.processor = processor
        self.files = files
        self.replacements = replacements
        self.output_folder = output_folder

    def run(self):
        """线程执行方法，遍历文件进行处理"""
        total_files = len(self.files)
        success_count = 0

        for i, file_path in enumerate(self.files):
            file_name = Path(file_path).name
            base_name = Path(file_path).stem

            self.progress_updated.emit(i, f"正在处理：{file_name} ({i + 1}/{total_files})")

            if self.output_folder:
                output_dir = Path(self.output_folder)
            else:
                file_dir = Path(file_path).parent
                output_dir = file_dir / "修改版"

            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f"{base_name}-修改版.docx"

            success = self.processor.process_document(
                file_path, self.replacements, str(output_path)
            )

            if success:
                success_count += 1

            progress = int(((i + 1) / total_files) * 100)
            self.progress_updated.emit(progress, "")

        self.processing_complete.emit(success_count, total_files)


class WriteNumberThread(QThread):
    """Excel编号写入线程"""
    progress_updated = pyqtSignal(int, str)
    write_complete = pyqtSignal(bool, str)

    def __init__(self, excel_processor, excel_file, output_path, number_text, 
                 use_fixed_mode, fixed_number, number_format):
        """初始化编号写入线程
        
        Args:
            excel_processor: ExcelProcessor实例
            excel_file: Excel文件路径
            output_path: 输出路径
            number_text: 编号文本
            use_fixed_mode: 是否固定模式
            fixed_number: 固定编号
            number_format: 编号格式
        """
        super().__init__()
        self.excel_processor = excel_processor
        self.excel_file = excel_file
        self.output_path = output_path
        self.number_text = number_text
        self.use_fixed_mode = use_fixed_mode
        self.fixed_number = fixed_number
        self.number_format = number_format

    def run(self):
        """线程执行方法"""
        self.progress_updated.emit(0, f"正在写入编号：{self.number_text}")

        success = self.excel_processor.write_number_to_excel(
            self.excel_file,
            self.output_path,
            self.number_text,
            use_fixed_mode=self.use_fixed_mode,
            fixed_number=self.fixed_number,
            number_format=self.number_format
        )

        self.progress_updated.emit(100, "")
        self.write_complete.emit(success, self.number_text)


class SuffixModifyThread(QThread):
    """文件后缀修改线程"""
    progress_updated = pyqtSignal(int, int, str)
    modify_complete = pyqtSignal(int, int)

    def __init__(self, processor, files, suffix):
        """初始化后缀修改线程
        
        Args:
            processor: SuffixProcessor实例
            files: 文件路径列表
            suffix: 要删除的后缀
        """
        super().__init__()
        self.processor = processor
        self.files = files
        self.suffix = suffix

    def run(self):
        """线程执行方法"""
        def progress_callback(current, total, filename):
            self.progress_updated.emit(current, total, filename)

        success, total = self.processor.remove_suffix_from_files(
            self.files, self.suffix, progress_callback
        )
        self.modify_complete.emit(success, total)


class ExportPdfThread(QThread):
    """PDF导出后台线程，避免阻塞UI和COM线程问题"""
    progress_updated = pyqtSignal(int, str)
    export_complete = pyqtSignal(bool, str)

    def __init__(self, excel_path, output_path):
        """初始化PDF导出线程
        
        Args:
            excel_path: 输入Excel文件路径
            output_path: 输出PDF文件路径
        """
        super().__init__()
        self.excel_path = excel_path
        self.output_path = output_path
        self._result = None
        self._error_message = ""

    def run(self):
        """线程执行方法：使用openpyxl生成PDF占位（实际由WordProcessor风格的COM操作）"""
        try:
            self.progress_updated.emit(10, "正在准备导出...")

            # 在线程中初始化COM
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()

            self.progress_updated.emit(30, "正在启动Excel...")

            excel = win32com.client.Dispatch("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            excel.ScreenUpdating = False

            self.progress_updated.emit(50, "正在打开文件...")

            wb = excel.Workbooks.Open(str(Path(self.excel_path).resolve()))

            self.progress_updated.emit(70, "正在设置页面...")

            for ws in wb.Worksheets:
                # 设置页面缩放：将所有列缩放为一页，行数不限
                ws.PageSetup.FitToPagesWide = 1
                ws.PageSetup.FitToPagesTall = 0
                ws.PageSetup.Zoom = False

                # 设置页边距（单位：厘米）
                ws.PageSetup.TopMargin = excel.CentimetersToPoints(1.3)
                ws.PageSetup.BottomMargin = excel.CentimetersToPoints(1.3)
                ws.PageSetup.LeftMargin = excel.CentimetersToPoints(0.9)
                ws.PageSetup.RightMargin = excel.CentimetersToPoints(0.9)

            # 选中所有工作表
            wb.Worksheets.Select()

            self.progress_updated.emit(85, "正在导出PDF...")

            # 导出为PDF (0 = xlTypePDF)
            wb.ExportAsFixedFormat(0, str(Path(self.output_path).resolve()))

            self.progress_updated.emit(95, "正在关闭Excel...")

            wb.Close(False)
            excel.Quit()

            pythoncom.CoUninitialize()

            self.progress_updated.emit(100, "导出完成")
            self.export_complete.emit(True, self.output_path)

        except ImportError as e:
            self._error_message = "缺少pywin32库，请安装：pip install pywin32"
            print(self._error_message)
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            self.export_complete.emit(False, self._error_message)
        except Exception as e:
            self._error_message = f"导出失败: {str(e)}"
            print(f"导出PDF失败 {self.excel_path}: {str(e)}")
            import traceback
            traceback.print_exc()
            try:
                excel.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            self.export_complete.emit(False, self._error_message)


# ============================================================================
# 主界面
# ============================================================================
class WordProcessorGUI(QMainWindow):
    """主窗口类，包含三个功能选项卡"""

    def __init__(self):
        """初始化主窗口"""
        super().__init__()
        self.setWindowTitle("Word 文本处理工具")
        self.setMinimumSize(900, 800)

        screen_geometry = QApplication.desktop().availableGeometry()
        window_width = min(1000, screen_geometry.width() - 100)
        window_height = min(900, screen_geometry.height() - 100)
        self.resize(window_width, window_height)

        # 初始化处理器（懒加载PDF处理器以加速启动）
        self.processor = WordProcessor()
        self.excel_processor = ExcelProcessor()
        self.suffix_processor = SuffixProcessor()
        self.pdf_processor = None  # 延迟初始化，避免启动时加载pywin32

        # 状态变量
        self.selected_files = []
        self.output_folder = ""
        self.suffix_selected_files = []

        # 加载设置
        self.settings_file = Path(__file__).parent / "settings.json"
        self.number_settings = self._load_number_settings()

        # 初始化界面
        self._init_styles()
        self._create_main_layout()

    def _init_styles(self):
        """初始化界面样式"""
        self.primary_color = "#14B8A6"
        self.primary_dark = "#0D9488"
        self.primary_light = "#F0FDFA"
        self.primary_hover = "#0F766E"
        self.title_color = "#0A7A6E"
        self.success_color = "#10B981"
        self.warning_color = "#F59E0B"
        self.error_color = "#EF4444"
        self.text_color = "#000000"
        self.label_color = "#808080"
        self.default_text_color = "#000000"
        self.bg_color = "#E0F2FE"
        self.card_bg = "#F8FAFC"
        self.input_bg = "#FFFFFF"
        self.border_color = "#CBD5E1"
        self.button_bg = "#E5E7EB"
        self.button_hover = "#D1D5DB"
        self.button_pressed = "#9CA3AF"

        font = QFont("Microsoft YaHei UI", 9)
        QApplication.setFont(font)

        self.setStyleSheet(f"""
            QMainWindow {{
                background-color: {self.bg_color};
            }}
            QWidget {{
                font-family: Microsoft YaHei UI;
                font-size: 9pt;
                color: {self.text_color};
            }}
            QTabWidget::pane {{
                border: 1px solid {self.border_color};
                background-color: {self.card_bg};
                border-radius: 6px;
            }}
            QTabBar::tab {{
                background-color: #DBEAFE;
                color: {self.text_color};
                padding: 8px 20px;
                margin-right: 2px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                min-width: 100px;
                font-size: 9pt;
            }}
            QTabBar::tab:selected {{
                background-color: {self.card_bg};
                color: {self.primary_color};
                font-weight: bold;
            }}
            QTabBar::tab:hover:!selected {{
                background-color: #BFDBFE;
            }}
            QGroupBox {{
                background-color: {self.card_bg};
                border: 1px solid {self.border_color};
                border-radius: 6px;
                margin-top: 8px;
                padding: 10px;
                font-weight: bold;
            }}
            QGroupBox::title {{
                color: {self.primary_color};
                font-weight: bold;
                padding-left: 6px;
                padding-right: 6px;
                subcontrol-origin: margin;
                subcontrol-position: top left;
            }}
            QPushButton {{
                background-color: {self.button_bg};
                color: {self.text_color};
                border: 1px solid {self.border_color};
                border-radius: 6px;
                padding: 8px 18px;
                font-weight: bold;
                font-size: 9pt;
                min-width: 85px;
            }}
            QPushButton:hover {{
                background-color: {self.button_hover};
            }}
            QPushButton:pressed {{
                background-color: {self.button_pressed};
            }}
            QPushButton:disabled {{
                background-color: #F3F4F6;
                color: #9CA3AF;
            }}
            QLineEdit {{
                border: 1px solid {self.border_color};
                border-radius: 6px;
                padding: 6px 10px;
                background-color: {self.input_bg};
                font-size: 9pt;
                color: {self.text_color};
            }}
            QLineEdit:focus {{
                border-color: {self.primary_color};
            }}
            QLineEdit:readonly {{
                background-color: #F1F5F9;
                color: {self.text_color};
            }}
            QTableWidget {{
                border: 1px solid {self.border_color};
                border-radius: 6px;
                background-color: {self.input_bg};
                gridline-color: {self.border_color};
                font-size: 9pt;
                color: {self.text_color};
            }}
            QTableWidget::item {{
                padding: 4px 6px;
                color: {self.text_color};
            }}
            QTableWidget::item:selected {{
                background-color: {self.primary_light};
                color: {self.text_color};
            }}
            QTableWidget::item:focus {{
                outline: none;
                border: none;
                background-color: {self.primary_light};
            }}
            QHeaderView::section {{
                background-color: #BFDBFE;
                color: {self.text_color};
                font-weight: bold;
                padding: 2px 6px;
                border: none;
                border-bottom: 1px solid {self.primary_color};
            }}
            QProgressBar {{
                border: none;
                border-radius: 4px;
                background-color: {self.border_color};
                height: 12px;
                text-align: center;
                color: {self.text_color};
            }}
            QProgressBar::chunk {{
                background: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:0, stop:0 {self.primary_color}, stop:1 {self.primary_dark});
                border-radius: 4px;
            }}
            QScrollBar:vertical {{
                width: 10px;
                background-color: #F1F5F9;
                border-radius: 5px;
                margin: 0px;
            }}
            QScrollBar::handle:vertical {{
                background-color: #94A3B8;
                border-radius: 5px;
                min-height: 30px;
            }}
            QScrollBar::handle:vertical:hover {{
                background-color: #64748B;
            }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
                height: 0px;
            }}
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {{
                background: none;
            }}
            QLabel {{
                color: {self.text_color};
            }}
        """)

    def _load_number_settings(self):
        """加载编号设置从配置文件"""
        try:
            if self.settings_file.exists():
                with open(self.settings_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception:
            pass

        return {
            "prefix": "HT合同编号",
            "current_number": 1,
            "number_format": "HT001-2026-XM001-01",
            "fixed_number": ""
        }

    def _save_number_settings(self):
        """保存编号设置到配置文件"""
        try:
            with open(self.settings_file, 'w', encoding='utf-8') as f:
                json.dump(self.number_settings, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存设置失败: {e}")

    def _generate_number(self):
        """生成下一个编号并递增计数器"""
        number = self.number_settings["current_number"]
        if number <= 99:
            number_str = str(number).zfill(2)
        else:
            number_str = str(number)

        number_format = self.number_settings["number_format"]
        full_number = number_format.replace("yy", number_str)

        self.number_settings["current_number"] = number + 1
        self._save_number_settings()

        return full_number

    def _create_main_layout(self):
        """创建主布局，包含标题和选项卡"""
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(10, 8, 10, 8)
        main_layout.setSpacing(6)

        title_label = QLabel("Word 文本处理工具")
        title_label.setStyleSheet(f"font-size: 14pt; font-weight: bold; color: {self.title_color};")
        title_label.setAlignment(Qt.AlignCenter)
        main_layout.addWidget(title_label)

        desc_label = QLabel("批量替换 Word 文档中的文本，支持正文、表格、页眉页脚")
        desc_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        desc_label.setAlignment(Qt.AlignCenter)
        main_layout.addWidget(desc_label)

        # 创建选项卡
        self.tab_widget = QTabWidget()
        self.tab_widget.setStyleSheet("QTabWidget::tab-bar { alignment: center; }")

        # 三个选项卡
        self.text_replace_tab = QWidget()
        self.number_record_tab = QWidget()
        self.suffix_modify_tab = QWidget()

        self.tab_widget.addTab(self.text_replace_tab, "文本替换")
        self.tab_widget.addTab(self.number_record_tab, "记录编号")
        self.tab_widget.addTab(self.suffix_modify_tab, "后缀修改")

        main_layout.addWidget(self.tab_widget)

        # 初始化各选项卡
        self._create_text_replace_tab()
        self._create_number_record_tab()
        self._create_suffix_modify_tab()

    # ========================================================================
    # 选项卡1：文本替换
    # ========================================================================
    def _create_text_replace_tab(self):
        """创建文本替换选项卡界面
        
        功能：批量替换Word文档中的文本
        支持：正文段落、表格内容、页眉页脚
        """
        layout = QVBoxLayout(self.text_replace_tab)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(6)

        # 文件选择区域
        self._create_file_section(layout)
        
        # 已选文件列表（可滚动）
        self._create_file_list_section(layout)
        
        # 操作按钮
        self._create_file_buttons(layout)
        
        # 替换规则设置
        self._create_replacement_section(layout)
        
        # 保存设置
        self._create_output_section(layout)
        
        # 进度显示
        self._create_progress_section(layout)
        
        # 执行按钮
        self._create_button_section(layout)

        layout.addStretch()

    def _create_file_section(self, parent_layout):
        """创建文件选择区域
        
        包含：文件路径显示、浏览文件/文件夹按钮、文件统计信息
        """
        group_box = QGroupBox("文件选择（可选择文件夹或文件）")
        group_layout = QVBoxLayout(group_box)
        group_layout.setSpacing(4)

        browse_layout = QHBoxLayout()
        browse_layout.setSpacing(5)

        browse_layout.addWidget(QLabel("选择文件/文件夹:"))

        self.file_path_edit = QLineEdit()
        self.file_path_edit.setReadOnly(True)
        self.file_path_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        browse_layout.addWidget(self.file_path_edit)

        browse_file_btn = QPushButton("浏览文件")
        browse_file_btn.clicked.connect(self._browse_files)
        browse_layout.addWidget(browse_file_btn)

        browse_folder_btn = QPushButton("浏览文件夹")
        browse_folder_btn.clicked.connect(self._browse_folder)
        browse_layout.addWidget(browse_folder_btn)

        group_layout.addLayout(browse_layout)

        info_layout = QHBoxLayout()
        info_layout.setSpacing(5)

        self.file_count_label = QLabel("未选择文件")
        self.file_count_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        info_layout.addWidget(self.file_count_label)

        self.file_path_detail_label = QLabel("")
        self.file_path_detail_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        self.file_path_detail_label.setWordWrap(True)
        info_layout.addWidget(self.file_path_detail_label)
        info_layout.addStretch()

        group_layout.addLayout(info_layout)

        parent_layout.addWidget(group_box)

    def _create_file_list_section(self, parent_layout):
        """创建已选文件列表区域
        
        优化：
        - 表格占满整个GroupBox区域
        - 支持垂直滚动，滚动条样式优化
        - 行高紧凑
        """
        group_box = QGroupBox("已选文件列表（双击可删除选中项）")
        group_layout = QVBoxLayout(group_box)
        group_layout.setContentsMargins(6, 12, 6, 6)
        group_layout.setSpacing(0)

        # 创建文件表格
        self.file_table = QTableWidget()
        self.file_table.setColumnCount(1)
        self.file_table.setHorizontalHeaderLabels(["文件名"])
        self.file_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.file_table.horizontalHeader().setFixedHeight(22)
        self.file_table.horizontalHeader().setStyleSheet("""
            QHeaderView::section {
                padding: 2px 6px;
                height: 22px;
                background-color: #BFDBFE;
                border: none;
                border-bottom: 1px solid #14B8A6;
            }
        """)
        self.file_table.verticalHeader().setVisible(False)
        self.file_table.verticalHeader().setDefaultSectionSize(20)
        self.file_table.verticalHeader().setMinimumSectionSize(18)
        self.file_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.file_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.file_table.doubleClicked.connect(self._remove_selected_file)
        self.file_table.setStyleSheet("""
            QTableWidget {
                border: 1px solid #CBD5E1;
                border-radius: 4px;
                background-color: #FFFFFF;
                font-size: 9pt;
            }
            QTableWidget::item {
                padding: 1px 6px;
                border: none;
            }
            QTableWidget::item:selected {
                background-color: #F0FDFA;
            }
        """)
        
        # 设置表格可滚动，占满区域
        self.file_table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.file_table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.file_table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.file_table.setMinimumHeight(132)

        group_layout.addWidget(self.file_table)

        parent_layout.addWidget(group_box)

    def _create_file_buttons(self, parent_layout):
        """创建文件操作按钮（添加/删除）"""
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(6)

        add_btn = QPushButton("+ 添加文件")
        add_btn.clicked.connect(self._add_file)
        add_btn.setFixedHeight(29)
        btn_layout.addWidget(add_btn)

        remove_btn = QPushButton("- 删除文件")
        remove_btn.clicked.connect(self._remove_selected_file)
        remove_btn.setFixedHeight(29)
        btn_layout.addWidget(remove_btn)

        btn_layout.addStretch()

        parent_layout.addLayout(btn_layout)

    def _create_replacement_section(self, parent_layout):
        """创建替换规则设置区域
        
        优化：
        - 表格占满整个区域
        - 编辑单元格时无边框间隔，内容完整显示
        - 固定行高，默认不显示滚动条
        - 重置按钮靠下方，与表格不重叠
        """
        group_box = QGroupBox("替换规则设置（双击单元格编辑）")
        group_layout = QVBoxLayout(group_box)
        group_layout.setContentsMargins(6, 12, 6, 6)
        group_layout.setSpacing(12)

        # 创建替换规则表格
        self.replacement_table = QTableWidget()
        self.replacement_table.setColumnCount(3)
        self.replacement_table.setHorizontalHeaderLabels(["类别", "查找", "替换"])
        
        # 设置列宽
        self.replacement_table.setColumnWidth(0, 100)
        self.replacement_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Fixed)
        self.replacement_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.replacement_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.replacement_table.horizontalHeader().setFixedHeight(24)
        self.replacement_table.horizontalHeader().setStyleSheet("""
            QHeaderView::section {
                padding: 2px 6px;
                height: 24px;
                background-color: #BFDBFE;
                border: none;
                border-bottom: 1px solid #14B8A6;
            }
        """)
        
        # 设置固定行高
        self.replacement_table.verticalHeader().setVisible(False)
        self.replacement_table.verticalHeader().setDefaultSectionSize(24)
        self.replacement_table.verticalHeader().setMinimumSectionSize(22)
        
        self.replacement_table.setSelectionBehavior(QAbstractItemView.SelectItems)
        self.replacement_table.cellDoubleClicked.connect(self._edit_cell)
        self.replacement_table.horizontalHeader().setMinimumSectionSize(80)
        
        # 设置表格样式：编辑时无边框，降低padding
        self.replacement_table.setStyleSheet("""
            QTableWidget {
                border: 1px solid #CBD5E1;
                border-radius: 4px;
                background-color: #FFFFFF;
                gridline-color: #E2E8F0;
            }
            QTableWidget::item {
                padding: 2px 6px;
                border: none;
            }
            QTableWidget::item:selected {
                background-color: #F0FDFA;
            }
            QTableWidget::item:focus {
                outline: none;
                border: none;
                background-color: #F0FDFA;
            }
            QLineEdit {
                border: 1px solid #CBD5E1;
                border-radius: 3px;
                padding: 2px 4px;
                background-color: #FFFFFF;
                margin: 0px;
            }
        """)
        
        # 固定高度，4行数据 + 表头，不显示滚动条
        self.replacement_table.setFixedHeight(130)
        self.replacement_table.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        group_layout.addWidget(self.replacement_table)

        # 重置按钮 - 靠下方，与表格有足够间距
        btn_layout = QHBoxLayout()
        btn_layout.setContentsMargins(0, 0, 0, 0)
        reset_btn = QPushButton("重置为默认值")
        reset_btn.clicked.connect(self._reset_to_defaults)
        reset_btn.setMinimumHeight(38)
        reset_btn.setMinimumWidth(130)
        reset_btn.setStyleSheet("""
            QPushButton {
                padding: 10px 20px;
                font-size: 9pt;
            }
        """)
        btn_layout.addWidget(reset_btn)
        btn_layout.addStretch()

        group_layout.addLayout(btn_layout)

        parent_layout.addWidget(group_box)

        # 初始化默认规则
        self._init_default_rules()

    def _create_output_section(self, parent_layout):
        """创建保存设置区域"""
        group_box = QGroupBox("保存设置")
        group_layout = QVBoxLayout(group_box)
        group_layout.setSpacing(4)

        browse_layout = QHBoxLayout()
        browse_layout.setSpacing(5)

        browse_layout.addWidget(QLabel("保存路径:"))

        self.output_path_edit = QLineEdit()
        self.output_path_edit.setReadOnly(True)
        self.output_path_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        browse_layout.addWidget(self.output_path_edit)

        browse_btn = QPushButton("浏览")
        browse_btn.clicked.connect(self._browse_output)
        browse_btn.setFixedHeight(29)
        browse_layout.addWidget(browse_btn)

        group_layout.addLayout(browse_layout)

        self.output_info_label = QLabel("未选择保存路径时将按规则自动创建文件夹，默认在同级目录下新建修改版文件夹")
        self.output_info_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        self.output_info_label.setWordWrap(True)
        group_layout.addWidget(self.output_info_label)

        parent_layout.addWidget(group_box)

    def _create_progress_section(self, parent_layout):
        """创建进度显示区域"""
        group_box = QGroupBox("处理进度")
        group_layout = QVBoxLayout(group_box)
        group_layout.setSpacing(4)

        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        group_layout.addWidget(self.progress_bar)

        self.progress_label = QLabel("就绪")
        self.progress_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        group_layout.addWidget(self.progress_label)

        parent_layout.addWidget(group_box)

    def _create_button_section(self, parent_layout):
        """创建执行按钮区域"""
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)
        btn_layout.setContentsMargins(0, 0, 0, 0)

        self.start_btn = QPushButton("开始处理")
        self.start_btn.clicked.connect(self._start_processing)
        self.start_btn.setFixedHeight(33)
        self.start_btn.setMinimumWidth(100)

        clear_btn = QPushButton("清空设置")
        clear_btn.clicked.connect(self._clear_all)
        clear_btn.setFixedHeight(33)
        clear_btn.setMinimumWidth(100)

        btn_layout.addStretch()
        btn_layout.addWidget(self.start_btn)
        btn_layout.addWidget(clear_btn)
        btn_layout.addStretch()

        parent_layout.addLayout(btn_layout)

    def _init_default_rules(self):
        """初始化默认替换规则"""
        self.default_rules = [
            ("项目编号", "HT001-2026-XM001-01", "HT001-2026-XM001-01"),
            ("公司名称", "", ""),
            ("系统名称", "", ""),
            ("人员名称", "", "")
        ]

        self.replacement_table.setRowCount(4)
        for row, (category, original, replacement) in enumerate(self.default_rules):
            # 类别列不可编辑
            category_item = QTableWidgetItem(category)
            category_item.setFlags(category_item.flags() & ~Qt.ItemIsEditable)
            category_item.setForeground(QColor(self.default_text_color))
            category_item.setTextAlignment(Qt.AlignCenter)
            self.replacement_table.setItem(row, 0, category_item)

            # 查找列
            original_item = QTableWidgetItem(original)
            if original:
                original_item.setForeground(QColor(self.default_text_color))
            self.replacement_table.setItem(row, 1, original_item)
            
            # 替换列
            self.replacement_table.setItem(row, 2, QTableWidgetItem(replacement))

    def _reset_to_defaults(self):
        """重置替换规则为默认值"""
        self.replacement_table.clearContents()
        self._init_default_rules()

    def _edit_cell(self, row, column):
        """双击编辑单元格（类别列不可编辑）"""
        if column == 0:
            return

        item = self.replacement_table.item(row, column)
        if item:
            self.replacement_table.editItem(item)

    def _browse_files(self):
        """浏览选择Word文件"""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "选择 Word 文件",
            "",
            "Word 文件 (*.docx *.doc);;所有文件 (*.*)"
        )
        if files:
            self.selected_files = list(files)
            self._update_file_display()

    def _browse_folder(self):
        """浏览选择包含Word文件的文件夹"""
        folder = QFileDialog.getExistingDirectory(self, "选择包含 Word 文件的文件夹")
        if folder:
            word_files = self.processor.get_word_files(folder)
            if word_files:
                self.selected_files = word_files
                self._update_file_display()
            else:
                QMessageBox.warning(self, "警告", "该文件夹下没有找到 Word 文件")

    def _add_file(self):
        """添加文件到列表"""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "添加 Word 文件",
            "",
            "Word 文件 (*.docx *.doc);;所有文件 (*.*)"
        )
        if files:
            for file in files:
                if file not in self.selected_files:
                    self.selected_files.append(file)
            self._update_file_display()

    def _remove_selected_file(self):
        """从列表中删除选中的文件"""
        selected_items = self.file_table.selectedItems()
        if selected_items:
            rows_to_remove = set()
            for item in selected_items:
                rows_to_remove.add(item.row())

            for row in sorted(rows_to_remove, reverse=True):
                file_item = self.file_table.item(row, 0)
                if file_item:
                    file_path = file_item.text()
                    if file_path in self.selected_files:
                        self.selected_files.remove(file_path)
                self.file_table.removeRow(row)

            self._update_file_display()

    def _update_file_display(self):
        """更新文件列表显示"""
        self.file_table.setRowCount(0)

        if self.selected_files:
            self.file_count_label.setText(f"已选择 {len(self.selected_files)} 个文件")
            self.file_count_label.setStyleSheet(f"color: {self.success_color}; font-weight: bold;")
            self.file_path_edit.setText(f"已选择 {len(self.selected_files)} 个文件")

            paths = "\n".join(self.selected_files[:5])
            if len(self.selected_files) > 5:
                paths += f"\n... 及其他 {len(self.selected_files) - 5} 个文件"
            self.file_path_detail_label.setText(paths)

            for file_path in self.selected_files:
                row = self.file_table.rowCount()
                self.file_table.insertRow(row)
                self.file_table.setItem(row, 0, QTableWidgetItem(file_path))
        else:
            self.file_count_label.setText("未选择文件")
            self.file_count_label.setStyleSheet(f"color: {self.label_color};")
            self.file_path_edit.setText("")
            self.file_path_detail_label.setText("")

    def _browse_output(self):
        """浏览选择输出文件夹"""
        folder = QFileDialog.getExistingDirectory(self, "选择保存路径")
        if folder:
            self.output_folder = folder
            self.output_path_edit.setText(folder)
            self.output_info_label.setText("已选择自定义保存路径")
            self.output_info_label.setStyleSheet(f"color: {self.success_color}; font-size: 9pt;")

    def _get_all_replacements(self):
        """获取所有替换规则"""
        replacements = {}
        for row in range(self.replacement_table.rowCount()):
            original_item = self.replacement_table.item(row, 1)
            replacement_item = self.replacement_table.item(row, 2)

            if original_item:
                original = original_item.text().strip()
                if original:
                    replacement = replacement_item.text().strip() if replacement_item else ""
                    replacements[original] = replacement

        return replacements

    def _clear_all(self):
        """清空所有设置"""
        self.selected_files = []
        self.output_folder = ""
        self.file_path_edit.setText("")
        self.output_path_edit.setText("")
        self.file_count_label.setText("未选择文件")
        self.file_count_label.setStyleSheet(f"color: {self.label_color};")
        self.output_info_label.setText("未选择保存路径时将按规则自动创建文件夹，默认在同级目录下新建修改版文件夹")
        self.output_info_label.setStyleSheet(f"color: {self.label_color}; font-size: 9pt;")
        self.progress_bar.setValue(0)
        self.progress_label.setText("就绪")

        self.file_table.setRowCount(0)
        self._reset_to_defaults()

    def _start_processing(self):
        """开始执行文本替换处理"""
        if not self.selected_files:
            QMessageBox.warning(self, "警告", "请先选择要处理的 Word 文件")
            return

        replacements = self._get_all_replacements()
        if not replacements:
            QMessageBox.warning(self, "警告", "请至少添加一条替换规则")
            return

        self.start_btn.setEnabled(False)
        self.progress_bar.setValue(0)
        self.progress_label.setText("准备处理...")

        self.processing_thread = ProcessingThread(
            self.processor,
            self.selected_files,
            replacements,
            self.output_folder
        )
        self.processing_thread.progress_updated.connect(self._update_progress)
        self.processing_thread.processing_complete.connect(self._processing_complete)
        self.processing_thread.start()

    def _update_progress(self, progress, message):
        """更新进度显示"""
        if message:
            self.progress_label.setText(message)
        if progress >= 0:
            self.progress_bar.setValue(progress)

    def _processing_complete(self, success_count, total_files):
        """处理完成回调"""
        self.start_btn.setEnabled(True)

        message = f"处理完成！\n成功：{success_count}/{total_files}"
        if success_count == total_files:
            message += "\n\n修改完成，已保存"
            QMessageBox.information(self, "完成", message)
        else:
            QMessageBox.warning(self, "完成", message)

        self.progress_label.setText("处理完成")

    # ========================================================================
    # 选项卡2：记录编号
    # ========================================================================
    def _create_number_record_tab(self):
        """创建记录编号选项卡界面
        
        功能：向Excel文件写入编号信息
        支持：自动编号模式、固定编号模式
        """
        layout = QVBoxLayout(self.number_record_tab)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(6)

        self._create_number_mode_section(layout)
        self._create_number_format_section(layout)
        self._create_number_fixed_section(layout)
        self._create_number_file_section(layout)
        self._create_number_progress_section(layout)
        self._create_number_button_section(layout)

        layout.addStretch()

    def _create_number_mode_section(self, parent_layout):
        """创建编号模式选择区域"""
        group_box = QGroupBox("编号模式")
        group_layout = QVBoxLayout(group_box)
        group_layout.setSpacing(6)

        self.number_mode = "auto"

        mode_layout = QHBoxLayout()
        mode_layout.setSpacing(15)

        # 自动编号模式按钮
        self.auto_radio = QPushButton("自动编号模式")
        self.auto_radio.setCheckable(True)
        self.auto_radio.setChecked(True)
        self.auto_radio.clicked.connect(lambda: self._on_mode_change("auto"))
        self.auto_radio.setFixedHeight(30)
        self.auto_radio.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.button_bg};
                border: 1px solid {self.border_color};
                border-radius: 6px;
                padding: 6px 14px;
                text-align: center;
                font-size: 9pt;
                color: {self.text_color};
            }}
            QPushButton:hover {{
                background-color: {self.button_hover};
            }}
            QPushButton:checked {{
                background-color: {self.primary_light};
                border-color: {self.primary_color};
                font-weight: bold;
            }}
        """)
        mode_layout.addWidget(self.auto_radio)

        # 固定编号模式按钮
        self.fixed_radio = QPushButton("固定编号模式")
        self.fixed_radio.setCheckable(True)
        self.fixed_radio.clicked.connect(lambda: self._on_mode_change("fixed"))
        self.fixed_radio.setFixedHeight(30)
        self.fixed_radio.setStyleSheet(f"""
            QPushButton {{
                background-color: {self.button_bg};
                border: 1px solid {self.border_color};
                border-radius: 6px;
                padding: 6px 14px;
                text-align: center;
                font-size: 9pt;
                color: {self.text_color};
            }}
            QPushButton:hover {{
                background-color: {self.button_hover};
            }}
            QPushButton:checked {{
                background-color: {self.primary_light};
                border-color: {self.primary_color};
                font-weight: bold;
            }}
        """)
        mode_layout.addWidget(self.fixed_radio)

        group_layout.addLayout(mode_layout)

        parent_layout.addWidget(group_box)

    def _create_number_format_section(self, parent_layout):
        """创建编号格式设置区域"""
        group_box = QGroupBox("编号格式设置")
        group_layout = QVBoxLayout(group_box)
        group_layout.setSpacing(4)

        format_layout = QHBoxLayout()
        format_layout.setSpacing(5)

        format_layout.addWidget(QLabel("编号格式:"))

        self.number_format_edit = QLineEdit()
        self.number_format_edit.setText(self.number_settings.get("number_format", "HT001-2026-XM001-yy"))
        self.number_format_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        format_layout.addWidget(self.number_format_edit)

        group_layout.addLayout(format_layout)

        desc_label = QLabel("说明：使用 yy 作为序号占位符，序号自动递增，超过99后自动变为3位数")
        desc_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        group_layout.addWidget(desc_label)

        parent_layout.addWidget(group_box)

    def _create_number_fixed_section(self, parent_layout):
        """创建固定编号设置区域"""
        self.fixed_number_group_box = QGroupBox("固定编号（仅在固定编号模式下使用）")
        group_layout = QVBoxLayout(self.fixed_number_group_box)
        group_layout.setSpacing(4)

        fixed_layout = QHBoxLayout()
        fixed_layout.setSpacing(5)

        fixed_layout.addWidget(QLabel("固定编号:"))

        self.fixed_number_edit = QLineEdit()
        self.fixed_number_edit.setText(self.number_settings.get("fixed_number", ""))
        self.fixed_number_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        fixed_layout.addWidget(self.fixed_number_edit)

        group_layout.addLayout(fixed_layout)

        info_label = QLabel("固定编号将写入 E1:F1 合并单元格，编号格式将写入 E2:F2 合并单元格")
        info_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        group_layout.addWidget(info_label)

        parent_layout.addWidget(self.fixed_number_group_box)

        self.fixed_number_group_box.setVisible(self.number_mode == "fixed")

    def _on_mode_change(self, mode):
        """切换编号模式"""
        self.number_mode = mode
        if mode == "auto":
            self.auto_radio.setChecked(True)
            self.fixed_radio.setChecked(False)
            self.fixed_number_group_box.setVisible(False)
        else:
            self.fixed_radio.setChecked(True)
            self.auto_radio.setChecked(False)
            self.fixed_number_group_box.setVisible(True)

    def _create_number_file_section(self, parent_layout):
        """创建Excel文件选择和保存设置区域"""
        # Excel文件选择
        group_box = QGroupBox("选择 Excel 文件")
        group_layout = QVBoxLayout(group_box)
        group_layout.setSpacing(4)

        browse_layout = QHBoxLayout()
        browse_layout.setSpacing(5)

        browse_layout.addWidget(QLabel("Excel 文件:"))

        self.excel_file_edit = QLineEdit()
        self.excel_file_edit.setReadOnly(True)
        self.excel_file_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        browse_layout.addWidget(self.excel_file_edit)

        browse_btn = QPushButton("浏览")
        browse_btn.clicked.connect(self._browse_excel_file)
        browse_btn.setFixedHeight(29)
        browse_layout.addWidget(browse_btn)

        group_layout.addLayout(browse_layout)

        self.excel_file_info_label = QLabel("未选择文件")
        self.excel_file_info_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        group_layout.addWidget(self.excel_file_info_label)

        parent_layout.addWidget(group_box)

        # 保存设置
        group_box2 = QGroupBox("保存设置")
        group_layout2 = QVBoxLayout(group_box2)
        group_layout2.setSpacing(4)

        save_layout = QHBoxLayout()
        save_layout.setSpacing(5)

        save_layout.addWidget(QLabel("保存路径:"))

        self.number_output_path_edit = QLineEdit()
        self.number_output_path_edit.setReadOnly(True)
        self.number_output_path_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        save_layout.addWidget(self.number_output_path_edit)

        save_browse_btn = QPushButton("浏览")
        save_browse_btn.clicked.connect(self._browse_number_output)
        save_browse_btn.setFixedHeight(29)
        save_layout.addWidget(save_browse_btn)

        group_layout2.addLayout(save_layout)

        self.number_output_info_label = QLabel("未选择保存路径时将自动保存在源文件同级目录，文件名后添加-修改版")
        self.number_output_info_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        self.number_output_info_label.setWordWrap(True)
        group_layout2.addWidget(self.number_output_info_label)

        parent_layout.addWidget(group_box2)

    def _create_number_progress_section(self, parent_layout):
        """创建编号处理进度显示区域"""
        group_box = QGroupBox("处理进度")
        group_layout = QVBoxLayout(group_box)
        group_layout.setSpacing(4)

        self.number_progress_bar = QProgressBar()
        self.number_progress_bar.setRange(0, 100)
        self.number_progress_bar.setValue(0)
        group_layout.addWidget(self.number_progress_bar)

        self.number_progress_label = QLabel("就绪")
        self.number_progress_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        group_layout.addWidget(self.number_progress_label)

        parent_layout.addWidget(group_box)

    def _create_number_button_section(self, parent_layout):
        """创建编号操作按钮区域
        
        包含三个按钮：
        - 写入编号：仅执行编号写入
        - 导出PDF：仅执行格式转换
        - 一键转换：先写入编号再导出PDF
        """
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)

        self.write_number_btn = QPushButton("写入编号")
        self.write_number_btn.clicked.connect(self._start_write_number)
        self.write_number_btn.setFixedHeight(33)
        self.write_number_btn.setMinimumWidth(100)

        reset_number_btn = QPushButton("重置编号")
        reset_number_btn.clicked.connect(self._reset_number)
        reset_number_btn.setFixedHeight(33)
        reset_number_btn.setMinimumWidth(100)

        self.export_pdf_btn = QPushButton("导出PDF")
        self.export_pdf_btn.clicked.connect(self._start_export_pdf)
        self.export_pdf_btn.setFixedHeight(33)
        self.export_pdf_btn.setMinimumWidth(100)

        self.one_click_convert_btn = QPushButton("一键转换")
        self.one_click_convert_btn.clicked.connect(self._start_one_click_convert)
        self.one_click_convert_btn.setFixedHeight(33)
        self.one_click_convert_btn.setMinimumWidth(100)

        btn_layout.addStretch()
        btn_layout.addWidget(self.write_number_btn)
        btn_layout.addWidget(reset_number_btn)
        btn_layout.addWidget(self.export_pdf_btn)
        btn_layout.addWidget(self.one_click_convert_btn)
        btn_layout.addStretch()

        parent_layout.addLayout(btn_layout)

    def _browse_excel_file(self):
        """浏览选择Excel文件"""
        file, _ = QFileDialog.getOpenFileName(
            self,
            "选择 Excel 文件",
            "",
            "Excel 文件 (*.xlsx *.xls);;所有文件 (*.*)"
        )
        if file:
            self.excel_file_path = file
            self.excel_file_edit.setText(file)
            self.excel_file_info_label.setText(f"已选择：{Path(file).name}")
            self.excel_file_info_label.setStyleSheet(f"color: {self.success_color};")
            self.number_output_path_edit.setText("")
            self.number_output_info_label.setText("未选择保存路径时将自动保存在源文件同级目录，文件名后添加-修改版")
            self.number_output_info_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")

    def _browse_number_output(self):
        """浏览选择编号输出路径"""
        folder = QFileDialog.getExistingDirectory(self, "选择保存路径")
        if folder:
            self.number_output_folder = folder
            self.number_output_path_edit.setText(folder)
            self.number_output_info_label.setText("已选择自定义保存路径")
            self.number_output_info_label.setStyleSheet(f"color: {self.success_color}; font-size: 8pt;")

    def _start_write_number(self):
        """开始写入编号"""
        if not hasattr(self, 'excel_file_path') or not self.excel_file_path:
            QMessageBox.warning(self, "警告", "请先选择要写入的 Excel 文件")
            return

        number_format = self.number_format_edit.text().strip()
        if not number_format:
            QMessageBox.warning(self, "警告", "请输入编号格式")
            return

        self.number_settings["number_format"] = number_format

        use_fixed_mode = self.number_mode == "fixed"
        fixed_number = ""

        if use_fixed_mode:
            fixed_number = self.fixed_number_edit.text().strip()
            if not fixed_number:
                QMessageBox.warning(self, "警告", "请输入固定编号")
                return
            self.number_settings["fixed_number"] = fixed_number

        self._save_number_settings()

        output_path = getattr(self, 'number_output_folder', None)
        if output_path:
            src_file = Path(self.excel_file_path)
            output_file = Path(output_path) / f"{src_file.stem}-修改版{src_file.suffix}"
        else:
            src_file = Path(self.excel_file_path)
            output_file = src_file.parent / f"{src_file.stem}-修改版{src_file.suffix}"

        self.write_number_btn.setEnabled(False)
        self.number_progress_label.setText("正在处理...")
        self.number_progress_bar.setValue(0)

        number_text = self._generate_number()

        self.write_thread = WriteNumberThread(
            self.excel_processor,
            self.excel_file_path,
            str(output_file),
            number_text,
            use_fixed_mode,
            fixed_number,
            number_format
        )
        self.write_thread.progress_updated.connect(self._update_number_progress)
        self.write_thread.write_complete.connect(self._write_number_complete)
        self.write_thread.start()

    def _update_number_progress(self, progress, message):
        """更新编号处理进度"""
        if message:
            self.number_progress_label.setText(message)
        if progress >= 0:
            self.number_progress_bar.setValue(progress)

    def _write_number_complete(self, success, number_text):
        """编号写入完成回调"""
        self.write_number_btn.setEnabled(True)

        use_fixed_mode = self.number_mode == "fixed"
        fixed_number = getattr(self, 'fixed_number_edit', None)
        fixed_number = fixed_number.text().strip() if fixed_number else ""

        if success:
            if use_fixed_mode:
                self.number_progress_label.setText(f"写入成功！固定编号：{fixed_number}，格式编号：{number_text}")
                QMessageBox.information(self, "完成", f"编号写入成功！\n\n固定编号（E1:F1）：{fixed_number}\n格式编号（E2:F2）：{number_text}")
            else:
                self.number_progress_label.setText(f"写入成功！编号：{number_text}")
                QMessageBox.information(self, "完成", f"编号写入成功！\n\n编号（E1:F1）：{number_text}")
        else:
            self.number_progress_label.setText("写入失败")
            self.number_progress_label.setStyleSheet(f"color: {self.error_color};")
            QMessageBox.critical(self, "错误", "编号写入失败，请检查文件是否被占用或格式是否正确")

    def _reset_number(self):
        """重置编号计数器"""
        reply = QMessageBox.question(
            self,
            "确认",
            "确定要重置编号吗？当前编号将恢复为 01",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            self.number_settings["current_number"] = 1
            self._save_number_settings()
            QMessageBox.information(self, "完成", "编号已重置为01")

    def _get_number_output_path(self):
        """获取编号写入的输出路径
        
        复用保存设置中的路径，未选择时自动保存在源文件同级目录
        
        Returns:
            str: 输出文件路径
        """
        output_folder = getattr(self, 'number_output_folder', None)
        if output_folder:
            src_file = Path(self.excel_file_path)
            return str(Path(output_folder) / f"{src_file.stem}-修改版.xlsx")
        else:
            src_file = Path(self.excel_file_path)
            return str(src_file.parent / f"{src_file.stem}-修改版.xlsx")

    def _get_pdf_output_path(self):
        """获取PDF导出的输出路径
        
        复用保存设置中的路径，未选择时自动保存在源文件同级目录
        
        Returns:
            str: 输出PDF文件路径
        """
        output_folder = getattr(self, 'number_output_folder', None)
        if output_folder:
            src_file = Path(self.excel_file_path)
            return str(Path(output_folder) / f"{src_file.stem}-修改版.pdf")
        else:
            src_file = Path(self.excel_file_path)
            return str(src_file.parent / f"{src_file.stem}-修改版.pdf")

    def _start_export_pdf(self):
        """开始导出PDF
        
        仅执行格式转换：
        - 选中所有工作表
        - 所有列缩放为一页
        - 上下页边距1.3cm，左右页边距0.9cm
        - 导出为PDF格式
        
        使用后台线程避免UI阻塞
        """
        if not hasattr(self, 'excel_file_path') or not self.excel_file_path:
            QMessageBox.warning(self, "警告", "请先选择Excel文件")
            return

        if not Path(self.excel_file_path).exists():
            QMessageBox.warning(self, "警告", "Excel文件不存在")
            return

        output_file = self._get_pdf_output_path()

        # 确保输出目录存在
        output_dir = Path(output_file).parent
        try:
            output_dir.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"无法创建输出目录：{str(e)}")
            return

        # 禁用所有相关按钮
        self.export_pdf_btn.setEnabled(False)
        self.write_number_btn.setEnabled(False)
        self.one_click_convert_btn.setEnabled(False)

        self.number_progress_label.setText("正在导出PDF...")
        self.number_progress_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        self.number_progress_bar.setValue(0)

        # 使用后台线程
        self.pdf_thread = ExportPdfThread(
            self.excel_file_path,
            output_file
        )
        self.pdf_thread.progress_updated.connect(self._update_number_progress)
        self.pdf_thread.export_complete.connect(self._on_export_pdf_complete)
        self.pdf_thread.start()

    def _on_export_pdf_complete(self, success, result):
        """PDF导出完成回调"""
        self.export_pdf_btn.setEnabled(True)
        self.write_number_btn.setEnabled(True)
        self.one_click_convert_btn.setEnabled(True)

        if success:
            self.number_progress_label.setText("PDF导出成功！")
            self.number_progress_label.setStyleSheet(f"color: {self.success_color}; font-size: 8pt;")
            self.number_progress_bar.setValue(100)
            QMessageBox.information(self, "完成", f"PDF导出成功！\n\n保存路径：{result}")
        else:
            self.number_progress_label.setText("PDF导出失败")
            self.number_progress_label.setStyleSheet(f"color: {self.error_color}; font-size: 8pt;")
            QMessageBox.critical(self, "错误", f"PDF导出失败：\n{result}\n\n请检查：\n1. Excel文件是否被占用\n2. 输出路径是否有写入权限\n3. 是否已安装Microsoft Excel\n4. pywin32库是否已安装")

    def _start_one_click_convert(self):
        """一键转换：先写入编号，再导出PDF
        
        流程：
        1. 按当前编号模式写入编号到Excel
        2. 将写入后的Excel导出为PDF
        """
        if not hasattr(self, 'excel_file_path') or not self.excel_file_path:
            QMessageBox.warning(self, "警告", "请先选择Excel文件")
            return

        number_format = self.number_format_edit.text().strip()
        if not number_format:
            QMessageBox.warning(self, "警告", "请输入编号格式")
            return

        # 先执行写入编号
        self.number_settings["number_format"] = number_format

        number_text = self._generate_number_text()

        use_fixed_mode = self.number_mode == "fixed"
        fixed_number = getattr(self, 'fixed_number_edit', None)
        fixed_number = fixed_number.text().strip() if fixed_number else ""

        self.write_number_btn.setEnabled(False)
        self.export_pdf_btn.setEnabled(False)
        self.one_click_convert_btn.setEnabled(False)
        self.number_progress_label.setText("正在写入编号...")
        self.number_progress_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        self.number_progress_bar.setValue(0)

        output_file = self._get_number_output_path()

        self.write_thread = WriteNumberThread(
            self.excel_processor,
            self.excel_file_path,
            output_file,
            number_text,
            use_fixed_mode,
            fixed_number,
            number_format
        )
        self.write_thread.progress_updated.connect(self._update_number_progress)
        self.write_thread.write_complete.connect(self._one_click_write_complete)
        self.write_thread.start()

    def _one_click_write_complete(self, success, number_text):
        """一键转换中编号写入完成后的回调，继续执行PDF导出
        
        使用单独的后台线程处理PDF导出，避免COM线程问题
        """
        if not success:
            self.write_number_btn.setEnabled(True)
            self.export_pdf_btn.setEnabled(True)
            self.one_click_convert_btn.setEnabled(True)
            self.number_progress_label.setText("编号写入失败，无法继续导出PDF")
            self.number_progress_label.setStyleSheet(f"color: {self.error_color}; font-size: 8pt;")
            QMessageBox.critical(self, "错误", "编号写入失败，无法继续导出PDF")
            return

        # 写入成功，启动PDF导出后台线程
        # 使用写入后的xlsx文件作为源
        xlsx_output = self._get_number_output_path()
        pdf_output = self._get_pdf_output_path()

        # 确保输出目录存在
        output_dir = Path(pdf_output).parent
        try:
            output_dir.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            self.write_number_btn.setEnabled(True)
            self.export_pdf_btn.setEnabled(True)
            self.one_click_convert_btn.setEnabled(True)
            self.number_progress_label.setText("编号写入成功，但PDF输出目录创建失败")
            self.number_progress_label.setStyleSheet(f"color: {self.error_color}; font-size: 8pt;")
            QMessageBox.warning(self, "完成", f"编号写入成功！\n\n编号：{number_text}\n\n但PDF输出目录创建失败：{str(e)}")
            return

        self.number_progress_label.setText("正在导出PDF...")

        # 保存上下文以便回调使用
        self._one_click_number_text = number_text
        self._one_click_xlsx_path = xlsx_output

        # 使用后台线程
        self.pdf_thread = ExportPdfThread(
            xlsx_output,
            pdf_output
        )
        self.pdf_thread.progress_updated.connect(self._update_number_progress)
        self.pdf_thread.export_complete.connect(self._on_one_click_pdf_complete)
        self.pdf_thread.start()

    def _on_one_click_pdf_complete(self, success, result):
        """一键转换中PDF导出完成回调"""
        self.write_number_btn.setEnabled(True)
        self.export_pdf_btn.setEnabled(True)
        self.one_click_convert_btn.setEnabled(True)

        number_text = getattr(self, '_one_click_number_text', '')

        if success:
            self.number_progress_label.setText("一键转换完成！")
            self.number_progress_label.setStyleSheet(f"color: {self.success_color}; font-size: 8pt;")
            self.number_progress_bar.setValue(100)
            QMessageBox.information(self, "完成", f"一键转换完成！\n\n编号写入成功：{number_text}\nPDF导出成功\n\nPDF保存路径：{result}")
        else:
            self.number_progress_label.setText("编号写入成功，但PDF导出失败")
            self.number_progress_label.setStyleSheet(f"color: {self.error_color}; font-size: 8pt;")
            QMessageBox.warning(self, "完成", f"编号写入成功！\n\n编号：{number_text}\n\n但PDF导出失败：\n{result}\n\n请检查：\n1. Excel文件是否被占用\n2. 输出路径是否有写入权限\n3. 是否已安装Microsoft Excel")

    # ========================================================================
    # 选项卡3：后缀修改
    # ========================================================================
    def _create_suffix_modify_tab(self):
        """创建后缀修改选项卡界面
        
        功能：批量删除文件名中的"-修改版"后缀
        支持：选择文件或文件夹
        """
        layout = QVBoxLayout(self.suffix_modify_tab)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(6)

        # 文件选择区域
        self._create_suffix_file_section(layout)
        
        # 已选文件列表
        self._create_suffix_file_list_section(layout)
        
        # 操作按钮
        self._create_suffix_buttons(layout)
        
        # 进度显示
        self._create_suffix_progress_section(layout)
        
        # 执行按钮
        self._create_suffix_execute_section(layout)

        layout.addStretch()

    def _create_suffix_file_section(self, parent_layout):
        """创建后缀修改的文件选择区域"""
        group_box = QGroupBox("文件选择（可选择文件夹或文件）")
        group_layout = QVBoxLayout(group_box)
        group_layout.setSpacing(4)

        browse_layout = QHBoxLayout()
        browse_layout.setSpacing(5)

        browse_layout.addWidget(QLabel("选择文件/文件夹:"))

        self.suffix_file_path_edit = QLineEdit()
        self.suffix_file_path_edit.setReadOnly(True)
        self.suffix_file_path_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        browse_layout.addWidget(self.suffix_file_path_edit)

        browse_file_btn = QPushButton("浏览文件")
        browse_file_btn.clicked.connect(self._suffix_browse_files)
        browse_layout.addWidget(browse_file_btn)

        browse_folder_btn = QPushButton("浏览文件夹")
        browse_folder_btn.clicked.connect(self._suffix_browse_folder)
        browse_layout.addWidget(browse_folder_btn)

        group_layout.addLayout(browse_layout)

        info_layout = QHBoxLayout()
        info_layout.setSpacing(5)

        self.suffix_file_count_label = QLabel("未选择文件")
        self.suffix_file_count_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        info_layout.addWidget(self.suffix_file_count_label)

        self.suffix_file_path_detail_label = QLabel("")
        self.suffix_file_path_detail_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        self.suffix_file_path_detail_label.setWordWrap(True)
        info_layout.addWidget(self.suffix_file_path_detail_label)
        info_layout.addStretch()

        group_layout.addLayout(info_layout)

        parent_layout.addWidget(group_box)

    def _create_suffix_file_list_section(self, parent_layout):
        """创建后缀修改的文件列表区域"""
        group_box = QGroupBox("已选文件列表（双击可删除选中项）")
        group_layout = QVBoxLayout(group_box)
        group_layout.setContentsMargins(6, 12, 6, 6)
        group_layout.setSpacing(0)

        self.suffix_file_table = QTableWidget()
        self.suffix_file_table.setColumnCount(2)
        self.suffix_file_table.setHorizontalHeaderLabels(["原文件名", "修改后"])
        self.suffix_file_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.suffix_file_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.suffix_file_table.horizontalHeader().setFixedHeight(22)
        self.suffix_file_table.horizontalHeader().setStyleSheet("""
            QHeaderView::section {
                padding: 2px 6px;
                height: 22px;
                background-color: #BFDBFE;
                border: none;
                border-bottom: 1px solid #14B8A6;
            }
        """)
        self.suffix_file_table.verticalHeader().setVisible(False)
        self.suffix_file_table.verticalHeader().setDefaultSectionSize(20)
        self.suffix_file_table.verticalHeader().setMinimumSectionSize(18)
        self.suffix_file_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.suffix_file_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.suffix_file_table.doubleClicked.connect(self._suffix_remove_selected_file)
        self.suffix_file_table.setStyleSheet("""
            QTableWidget {
                border: 1px solid #CBD5E1;
                border-radius: 4px;
                background-color: #FFFFFF;
                font-size: 9pt;
            }
            QTableWidget::item {
                padding: 1px 6px;
                border: none;
            }
            QTableWidget::item:selected {
                background-color: #F0FDFA;
            }
        """)
        self.suffix_file_table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.suffix_file_table.setMinimumHeight(132)

        group_layout.addWidget(self.suffix_file_table)

        parent_layout.addWidget(group_box)

    def _create_suffix_buttons(self, parent_layout):
        """创建后缀修改的文件操作按钮"""
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(6)

        add_btn = QPushButton("+ 添加文件")
        add_btn.clicked.connect(self._suffix_add_file)
        add_btn.setFixedHeight(29)
        btn_layout.addWidget(add_btn)

        remove_btn = QPushButton("- 删除文件")
        remove_btn.clicked.connect(self._suffix_remove_selected_file)
        remove_btn.setFixedHeight(29)
        btn_layout.addWidget(remove_btn)

        btn_layout.addStretch()

        parent_layout.addLayout(btn_layout)

    def _create_suffix_progress_section(self, parent_layout):
        """创建后缀修改的进度显示区域"""
        group_box = QGroupBox("处理进度")
        group_layout = QVBoxLayout(group_box)
        group_layout.setSpacing(4)

        self.suffix_progress_bar = QProgressBar()
        self.suffix_progress_bar.setRange(0, 100)
        self.suffix_progress_bar.setValue(0)
        group_layout.addWidget(self.suffix_progress_bar)

        self.suffix_progress_label = QLabel("就绪")
        self.suffix_progress_label.setStyleSheet(f"color: {self.label_color}; font-size: 8pt;")
        group_layout.addWidget(self.suffix_progress_label)

        parent_layout.addWidget(group_box)

    def _create_suffix_execute_section(self, parent_layout):
        """创建后缀修改的执行按钮区域"""
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)
        btn_layout.setContentsMargins(0, 0, 0, 0)

        self.suffix_start_btn = QPushButton("开始修改")
        self.suffix_start_btn.clicked.connect(self._suffix_start_processing)
        self.suffix_start_btn.setFixedHeight(33)
        self.suffix_start_btn.setMinimumWidth(100)

        clear_btn = QPushButton("清空列表")
        clear_btn.clicked.connect(self._suffix_clear_all)
        clear_btn.setFixedHeight(33)
        clear_btn.setMinimumWidth(100)

        btn_layout.addStretch()
        btn_layout.addWidget(self.suffix_start_btn)
        btn_layout.addWidget(clear_btn)
        btn_layout.addStretch()

        parent_layout.addLayout(btn_layout)

    def _suffix_browse_files(self):
        """浏览选择文件（后缀修改）"""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "选择文件",
            "",
            "所有文件 (*.*)"
        )
        if files:
            self.suffix_selected_files = list(files)
            self._suffix_update_file_display()

    def _suffix_browse_folder(self):
        """浏览选择文件夹（后缀修改）"""
        folder = QFileDialog.getExistingDirectory(self, "选择包含文件的文件夹")
        if folder:
            files = []
            folder_path = Path(folder)
            for ext in ['*.doc', '*.docx', '*.pdf', '*.xlsx', '*.xls', '*.txt']:
                for file in folder_path.glob(ext):
                    files.append(str(file))
            if files:
                self.suffix_selected_files = files
                self._suffix_update_file_display()
            else:
                QMessageBox.warning(self, "警告", "该文件夹下没有找到支持的文件")

    def _suffix_add_file(self):
        """添加文件到列表（后缀修改）"""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "添加文件",
            "",
            "所有文件 (*.*)"
        )
        if files:
            for file in files:
                if file not in self.suffix_selected_files:
                    self.suffix_selected_files.append(file)
            self._suffix_update_file_display()

    def _suffix_remove_selected_file(self):
        """从列表中删除选中的文件（后缀修改）"""
        selected_items = self.suffix_file_table.selectedItems()
        if selected_items:
            rows_to_remove = set()
            for item in selected_items:
                rows_to_remove.add(item.row())

            for row in sorted(rows_to_remove, reverse=True):
                file_item = self.suffix_file_table.item(row, 0)
                if file_item:
                    # 从显示的文件名还原完整路径
                    file_name = file_item.text()
                    for file_path in self.suffix_selected_files:
                        if Path(file_path).name == file_name:
                            self.suffix_selected_files.remove(file_path)
                            break
                self.suffix_file_table.removeRow(row)

            self._suffix_update_file_display()

    def _suffix_update_file_display(self):
        """更新文件列表显示（后缀修改）"""
        self.suffix_file_table.setRowCount(0)

        if self.suffix_selected_files:
            self.suffix_file_count_label.setText(f"已选择 {len(self.suffix_selected_files)} 个文件")
            self.suffix_file_count_label.setStyleSheet(f"color: {self.success_color}; font-weight: bold;")
            self.suffix_file_path_edit.setText(f"已选择 {len(self.suffix_selected_files)} 个文件")

            paths = "\n".join(self.suffix_selected_files[:5])
            if len(self.suffix_selected_files) > 5:
                paths += f"\n... 及其他 {len(self.suffix_selected_files) - 5} 个文件"
            self.suffix_file_path_detail_label.setText(paths)

            suffix = "-修改版"
            for file_path in self.suffix_selected_files:
                row = self.suffix_file_table.rowCount()
                self.suffix_file_table.insertRow(row)
                
                path = Path(file_path)
                original_name = path.name
                
                # 计算修改后的文件名（删除所有匹配后缀）
                new_stem = path.stem.replace(suffix, "")
                
                if new_stem and new_stem != path.stem:
                    new_name = new_stem + path.suffix
                else:
                    new_name = original_name + " (无需修改)"
                
                self.suffix_file_table.setItem(row, 0, QTableWidgetItem(original_name))
                self.suffix_file_table.setItem(row, 1, QTableWidgetItem(new_name))
        else:
            self.suffix_file_count_label.setText("未选择文件")
            self.suffix_file_count_label.setStyleSheet(f"color: {self.label_color};")
            self.suffix_file_path_edit.setText("")
            self.suffix_file_path_detail_label.setText("")

    def _suffix_clear_all(self):
        """清空后缀修改的所有设置"""
        self.suffix_selected_files = []
        self.suffix_file_path_edit.setText("")
        self.suffix_file_count_label.setText("未选择文件")
        self.suffix_file_count_label.setStyleSheet(f"color: {self.label_color};")
        self.suffix_file_path_detail_label.setText("")
        self.suffix_progress_bar.setValue(0)
        self.suffix_progress_label.setText("就绪")
        self.suffix_file_table.setRowCount(0)

    def _suffix_start_processing(self):
        """开始执行后缀修改"""
        if not self.suffix_selected_files:
            QMessageBox.warning(self, "警告", "请先选择要处理的文件")
            return

        # 检查是否有需要修改的文件
        suffix = "-修改版"
        files_to_modify = [f for f in self.suffix_selected_files if Path(f).stem.endswith(suffix)]
        
        if not files_to_modify:
            QMessageBox.warning(self, "警告", "选中的文件中没有包含'-修改版'后缀的文件")
            return

        reply = QMessageBox.question(
            self,
            "确认",
            f"确定要修改 {len(files_to_modify)} 个文件的名称吗？\n将删除文件名中的'-修改版'后缀。",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply != QMessageBox.Yes:
            return

        self.suffix_start_btn.setEnabled(False)
        self.suffix_progress_bar.setValue(0)
        self.suffix_progress_label.setText("准备处理...")

        self.suffix_thread = SuffixModifyThread(
            self.suffix_processor,
            self.suffix_selected_files,
            suffix
        )
        self.suffix_thread.progress_updated.connect(self._suffix_update_progress)
        self.suffix_thread.modify_complete.connect(self._suffix_modify_complete)
        self.suffix_thread.start()

    def _suffix_update_progress(self, current, total, filename):
        """更新后缀修改进度"""
        self.suffix_progress_label.setText(f"正在处理：{filename} ({current + 1}/{total})")
        progress = int(((current + 1) / total) * 100)
        self.suffix_progress_bar.setValue(progress)

    def _suffix_modify_complete(self, success_count, total_count):
        """后缀修改完成回调"""
        self.suffix_start_btn.setEnabled(True)
        
        message = f"处理完成！\n成功修改：{success_count} 个文件"
        QMessageBox.information(self, "完成", message)
        self.suffix_progress_label.setText("处理完成")
        
        # 刷新显示
        self._suffix_update_file_display()


def main():
    """程序入口函数"""
    app = QApplication(sys.argv)

    font = QFont("Microsoft YaHei UI", 10)
    app.setFont(font)

    window = WordProcessorGUI()
    window.show()

    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
